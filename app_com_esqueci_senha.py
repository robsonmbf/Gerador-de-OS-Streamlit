import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile
from io import BytesIO
import time
import re
from datetime import datetime, timedelta
import os
import hashlib
import secrets
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# Configuração da página
st.set_page_config(
    page_title="Gerador de OS Profissional",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# CSS personalizado
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 2rem;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 2rem;
    }
    .info-card {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 4px solid #667eea;
        margin: 1rem 0;
    }
    .success-card {
        background: #d4edda;
        color: #155724;
        padding: 1.5rem;
        border-radius: 10px;
        border: 1px solid #c3e6cb;
        margin: 1rem 0;
    }
    .warning-card {
        background: #fff3cd;
        color: #856404;
        padding: 1.5rem;
        border-radius: 10px;
        border: 1px solid #ffeaa7;
        margin: 1rem 0;
    }
    .reset-card {
        background: #e2e3e5;
        color: #383d41;
        padding: 1.5rem;
        border-radius: 10px;
        border: 1px solid #d6d8db;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Sistema de usuários simulado (em produção usar banco de dados)
if 'users_db' not in st.session_state:
    st.session_state.users_db = {
        'admin': {
            'password_hash': hashlib.sha256('admin123'.encode()).hexdigest(),
            'email': 'admin@sistema.com',
            'full_name': 'Administrador',
            'is_admin': True,
            'credits': 999999,
            'created_at': datetime.now()
        }
    }

if 'reset_tokens' not in st.session_state:
    st.session_state.reset_tokens = {}

# Configurações de email (em produção usar variáveis de ambiente)
EMAIL_CONFIG = {
    'smtp_server': 'smtp.gmail.com',
    'smtp_port': 587,
    'email': 'seu-sistema@gmail.com',  # ← CONFIGURAR
    'password': 'sua-senha-app'        # ← CONFIGURAR
}

class AuthManager:
    def __init__(self):
        pass

    def hash_password(self, password):
        """Cria hash da senha"""
        return hashlib.sha256(password.encode()).hexdigest()

    def register_user(self, username, email, password, full_name=""):
        """Registrar novo usuário"""
        try:
            if username in st.session_state.users_db:
                return {"success": False, "message": "Usuário já existe"}

            # Verificar se email já existe
            for user, data in st.session_state.users_db.items():
                if data['email'] == email:
                    return {"success": False, "message": "Email já cadastrado"}

            if len(username) < 3:
                return {"success": False, "message": "Nome de usuário muito curto"}

            if len(password) < 6:
                return {"success": False, "message": "Senha deve ter pelo menos 6 caracteres"}

            if '@' not in email:
                return {"success": False, "message": "Email inválido"}

            # Criar usuário
            st.session_state.users_db[username] = {
                'password_hash': self.hash_password(password),
                'email': email,
                'full_name': full_name,
                'is_admin': False,
                'credits': 5,  # Créditos iniciais
                'created_at': datetime.now()
            }

            return {"success": True, "message": "Cadastro realizado! Você ganhou 5 créditos gratuitos."}

        except Exception as e:
            return {"success": False, "message": f"Erro interno: {str(e)}"}

    def login_user(self, username, password):
        """Fazer login"""
        try:
            if username not in st.session_state.users_db:
                return {"success": False, "message": "Usuário não encontrado"}

            user_data = st.session_state.users_db[username]

            if user_data['password_hash'] != self.hash_password(password):
                return {"success": False, "message": "Senha incorreta"}

            return {"success": True, "user": {
                'username': username,
                'email': user_data['email'],
                'full_name': user_data['full_name'],
                'is_admin': user_data['is_admin'],
                'credits': user_data['credits']
            }}

        except Exception as e:
            return {"success": False, "message": f"Erro: {str(e)}"}

    def generate_reset_token(self, email):
        """Gerar token de reset de senha"""
        try:
            # Verificar se email existe
            user_found = None
            for username, data in st.session_state.users_db.items():
                if data['email'] == email:
                    user_found = username
                    break

            if not user_found:
                return {"success": False, "message": "Email não encontrado"}

            # Gerar token único
            token = secrets.token_urlsafe(32)

            # Salvar token com expiração (1 hora)
            st.session_state.reset_tokens[token] = {
                'username': user_found,
                'email': email,
                'expires_at': datetime.now() + timedelta(hours=1)
            }

            # Enviar email (simulado)
            success = self.send_reset_email(email, token)

            if success:
                return {"success": True, "message": "Email de recuperação enviado!", "token": token}
            else:
                return {"success": True, "message": "Token gerado (Email simulado)", "token": token}

        except Exception as e:
            return {"success": False, "message": f"Erro: {str(e)}"}

    def send_reset_email(self, email, token):
        """Enviar email de reset (versão simplificada)"""
        try:
            # Em produção, configure SMTP real
            # Por agora, apenas simula o envio

            subject = "Recuperação de Senha - Gerador de OS"

            body = f"""
            Olá!

            Você solicitou a recuperação de senha do seu Gerador de OS.

            Use este token para redefinir sua senha:

            TOKEN: {token}

            Este token expira em 1 hora.

            Se você não solicitou esta recuperação, ignore este email.

            Atenciosamente,
            Sistema Gerador de OS
            """

            # Aqui você configuraria o envio real do email
            # Por enquanto, retorna True (simulado)
            return True

        except Exception as e:
            return False

    def reset_password(self, token, new_password):
        """Redefinir senha usando token"""
        try:
            if token not in st.session_state.reset_tokens:
                return {"success": False, "message": "Token inválido"}

            token_data = st.session_state.reset_tokens[token]

            # Verificar se token expirou
            if datetime.now() > token_data['expires_at']:
                del st.session_state.reset_tokens[token]
                return {"success": False, "message": "Token expirado"}

            if len(new_password) < 6:
                return {"success": False, "message": "Senha deve ter pelo menos 6 caracteres"}

            # Atualizar senha
            username = token_data['username']
            st.session_state.users_db[username]['password_hash'] = self.hash_password(new_password)

            # Remover token usado
            del st.session_state.reset_tokens[token]

            return {"success": True, "message": "Senha redefinida com sucesso!"}

        except Exception as e:
            return {"success": False, "message": f"Erro: {str(e)}"}

# Inicializar gerenciador de autenticação
auth_manager = AuthManager()

# Sistema de sessão
def init_session():
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'user' not in st.session_state:
        st.session_state.user = None
    if 'page' not in st.session_state:
        st.session_state.page = 'login'

def logout():
    st.session_state.logged_in = False
    st.session_state.user = None
    st.session_state.page = 'login'
    st.rerun()

# Função para gerar OS (mantém a funcionalidade original)
def gerar_os_do_funcionario(funcionario, modelo_docx):
    """Gera OS individual para um funcionário"""
    doc = Document(modelo_docx)

    # Dicionário de substituições
    substituicoes = {
        '{NOME_FUNCIONARIO}': str(funcionario.get('NOME', '')),
        '{FUNCAO}': str(funcionario.get('FUNCAO', '')),
        '{SETOR}': str(funcionario.get('SETOR', '')),
        '{DATA_AVALIACAO}': str(funcionario.get('DATA_AVALIACAO', '')),
        '{AGENTE_DE_RISCO}': str(funcionario.get('AGENTE_DE_RISCO', '')),
        '{FONTE_GERADORA}': str(funcionario.get('FONTE_GERADORA', '')),
        '{INTENSIDADE_CONCENTRACAO}': str(funcionario.get('INTENSIDADE_CONCENTRACAO', '')),
        '{UNIDADE_DE_MEDIDA}': str(funcionario.get('UNIDADE_DE_MEDIDA', '')),
        '{LIMITE_DE_TOLERANCIA}': str(funcionario.get('LIMITE_DE_TOLERANCIA', '')),
        '{TECNICA_UTILIZADA}': str(funcionario.get('TECNICA_UTILIZADA', '')),
        '{EPC_EXISTENTE}': str(funcionario.get('EPC_EXISTENTE', '')),
        '{EPC_RECOMENDADO}': str(funcionario.get('EPC_RECOMENDADO', '')),
        '{EPI_EXISTENTE}': str(funcionario.get('EPI_EXISTENTE', '')),
        '{EPI_RECOMENDADO}': str(funcionario.get('EPI_RECOMENDADO', '')),
        '{OBSERVACOES}': str(funcionario.get('OBSERVACOES', ''))
    }

    # Substituir nos parágrafos
    for paragraph in doc.paragraphs:
        for placeholder, value in substituicoes.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Substituir nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in substituicoes.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    return doc

def processar_os_lote(df_funcionarios, modelo_docx):
    """Processa lote de funcionários e gera todas as OS"""
    documentos_gerados = []

    progress_bar = st.progress(0)
    status_text = st.empty()

    for index, funcionario in df_funcionarios.iterrows():
        try:
            status_text.text(f"Gerando OS para: {funcionario.get('NOME', 'Funcionário')} ({index+1}/{len(df_funcionarios)})")

            doc_gerado = gerar_os_do_funcionario(funcionario, modelo_docx)

            # Gerar nome do arquivo
            nome_funcionario = str(funcionario.get('NOME', f'Funcionario_{index}')).strip()
            nome_funcionario = re.sub(r'[^\w\s-]', '', nome_funcionario).strip()
            nome_funcionario = re.sub(r'[-\s]+', '_', nome_funcionario)

            if not nome_funcionario:
                nome_funcionario = f'Funcionario_{index}'

            nome_arquivo = f"OS_{nome_funcionario}.docx"

            # Salvar documento em bytes
            doc_bytes = BytesIO()
            doc_gerado.save(doc_bytes)
            doc_bytes.seek(0)

            documentos_gerados.append({
                'nome': nome_arquivo,
                'conteudo': doc_bytes.getvalue()
            })

            progress_bar.progress((index + 1) / len(df_funcionarios))

        except Exception as e:
            st.error(f"Erro ao gerar OS para {funcionario.get('NOME', 'funcionário')}: {str(e)}")
            continue

    status_text.text("✅ Processamento concluído!")
    return documentos_gerados

def criar_zip_documentos(documentos):
    """Cria arquivo ZIP com todos os documentos"""
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for doc in documentos:
            zip_file.writestr(doc['nome'], doc['conteudo'])

    zip_buffer.seek(0)
    return zip_buffer.getvalue()

# Página de login com esqueci senha
def show_login_page():
    st.markdown("""
    <div class="main-header">
        <h1>🔐 Sistema Gerador de OS</h1>
        <p>Sistema Profissional com Recuperação de Senha</p>
        <small>👑 Admin: admin/admin123 | 💎 5 créditos grátis no cadastro</small>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3 = st.tabs(["🔑 Login", "📝 Cadastro", "🔄 Esqueci Senha"])

    with tab1:
        st.subheader("Entrar no Sistema")

        with st.form("login_form"):
            username = st.text_input("👤 Usuário")
            password = st.text_input("🔒 Senha", type="password")

            col1, col2 = st.columns([2, 1])

            with col1:
                login_btn = st.form_submit_button("🚀 Entrar", type="primary")

            with col2:
                forgot_btn = st.form_submit_button("❓ Esqueci senha", type="secondary")

            if login_btn and username and password:
                with st.spinner("Autenticando..."):
                    result = auth_manager.login_user(username, password)

                    if result['success']:
                        st.session_state.logged_in = True
                        st.session_state.user = result['user']
                        st.session_state.page = 'dashboard'
                        st.success("Login realizado com sucesso!")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error(result['message'])

            if forgot_btn:
                st.session_state.page = 'forgot_password'
                st.rerun()

            if not username or not password:
                if login_btn:
                    st.error("Preencha usuário e senha")

    with tab2:
        st.subheader("Criar Nova Conta")

        with st.form("register_form"):
            new_username = st.text_input("👤 Nome de usuário", help="Mínimo 3 caracteres")
            new_email = st.text_input("📧 Email")
            new_full_name = st.text_input("👨‍💼 Nome completo")
            new_password = st.text_input("🔒 Senha", type="password", help="Mínimo 6 caracteres")
            new_password_confirm = st.text_input("🔒 Confirmar senha", type="password")

            accept_terms = st.checkbox("✅ Aceito os termos de uso")

            if st.form_submit_button("📝 Criar Conta", type="primary"):
                if not all([new_username, new_email, new_password]):
                    st.error("❌ Preencha todos os campos obrigatórios")
                elif new_password != new_password_confirm:
                    st.error("❌ Senhas não coincidem")
                elif not accept_terms:
                    st.error("❌ Aceite os termos de uso")
                else:
                    with st.spinner("Criando conta..."):
                        result = auth_manager.register_user(new_username, new_email, new_password, new_full_name)

                        if result['success']:
                            st.success(result['message'])
                            st.info("🎉 Agora você pode fazer login!")
                        else:
                            st.error(result['message'])

    with tab3:
        st.subheader("🔄 Recuperar Senha")

        st.markdown("""
        <div class="info-card">
            <h4>📧 Como funciona:</h4>
            <p>1. Digite seu email cadastrado</p>
            <p>2. Receba um token de recuperação</p>
            <p>3. Use o token para criar nova senha</p>
            <p><small><strong>⚠️ O token expira em 1 hora</strong></small></p>
        </div>
        """, unsafe_allow_html=True)

        # Etapa 1: Solicitar reset
        if 'reset_step' not in st.session_state:
            st.session_state.reset_step = 1

        if st.session_state.reset_step == 1:
            with st.form("forgot_form"):
                email = st.text_input("📧 Digite seu email cadastrado")

                if st.form_submit_button("🔄 Enviar Token", type="primary"):
                    if email:
                        with st.spinner("Gerando token de recuperação..."):
                            result = auth_manager.generate_reset_token(email)

                            if result['success']:
                                st.success(result['message'])

                                # Mostrar token (em produção seria enviado por email)
                                st.markdown(f"""
                                <div class="warning-card">
                                    <h4>🔑 Seu Token de Recuperação:</h4>
                                    <code style="font-size: 1.2em; background: white; padding: 0.5rem; border-radius: 5px; display: block; margin: 0.5rem 0;">{result['token']}</code>
                                    <p><small><strong>💡 Em produção, este token seria enviado por email.</strong></small></p>
                                    <p><small>⏰ Expira em 1 hora</small></p>
                                </div>
                                """, unsafe_allow_html=True)

                                st.session_state.reset_step = 2
                                st.info("👇 Agora use o token abaixo para redefinir sua senha")

                            else:
                                st.error(result['message'])
                    else:
                        st.error("Digite um email")

        # Etapa 2: Redefinir senha
        if st.session_state.reset_step == 2:
            st.markdown("---")
            st.subheader("🔑 Redefinir Senha")

            with st.form("reset_form"):
                token = st.text_input("🔑 Token de Recuperação")
                new_password = st.text_input("🔒 Nova Senha", type="password", help="Mínimo 6 caracteres")
                confirm_password = st.text_input("🔒 Confirmar Nova Senha", type="password")

                col1, col2 = st.columns(2)

                with col1:
                    if st.form_submit_button("✅ Redefinir Senha", type="primary"):
                        if not all([token, new_password, confirm_password]):
                            st.error("❌ Preencha todos os campos")
                        elif new_password != confirm_password:
                            st.error("❌ Senhas não coincidem")
                        else:
                            with st.spinner("Redefinindo senha..."):
                                result = auth_manager.reset_password(token, new_password)

                                if result['success']:
                                    st.success(result['message'])
                                    st.balloons()
                                    st.info("🎉 Agora você pode fazer login com a nova senha!")

                                    # Reset do processo
                                    st.session_state.reset_step = 1
                                    time.sleep(2)
                                    st.rerun()
                                else:
                                    st.error(result['message'])

                with col2:
                    if st.form_submit_button("🔄 Solicitar Novo Token"):
                        st.session_state.reset_step = 1
                        st.rerun()

def show_dashboard():
    user = st.session_state.user

    # Header
    col1, col2, col3 = st.columns([2, 1, 1])

    with col1:
        st.markdown(f"""
        <div class="main-header">
            <h1>🏗️ Sistema Profissional</h1>
            <p>Bem-vindo, {user.get('full_name') or user['username']}!</p>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown(f"""
        <div class="info-card">
            <h3>💎 Créditos</h3>
            <h2>{user.get('credits', 0)}</h2>
            <small>disponíveis</small>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.write("")
        if st.button("🚪 Logout", type="secondary"):
            logout()

    # Menu
    st.markdown("---")
    col1, col2 = st.columns(2)

    with col1:
        if st.button("🏭 Gerar OS", type="primary", use_container_width=True):
            st.session_state.page = 'generator'
            st.rerun()

    with col2:
        if st.button("👤 Meu Perfil", type="secondary", use_container_width=True):
            st.session_state.page = 'profile'
            st.rerun()

def show_generator_page():
    user = st.session_state.user

    st.markdown(f"""
    <div class="main-header">
        <h1>🏭 Gerador de OS</h1>
        <p>💎 Você tem {user.get('credits', 0)} créditos disponíveis</p>
    </div>
    """, unsafe_allow_html=True)

    if user.get('credits', 0) <= 0:
        st.error("❌ Você não possui créditos suficientes!")
        st.info("💡 Contate o administrador para obter mais créditos.")

        if st.button("← Voltar ao Dashboard"):
            st.session_state.page = 'dashboard'
            st.rerun()
        return

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 📊 Planilha de Funcionários")
        arquivo_funcionarios = st.file_uploader(
            "Faça upload da planilha Excel (.xlsx)",
            type=['xlsx'],
            help="Planilha com dados dos funcionários"
        )

    with col2:
        st.markdown("#### 📄 Modelo de OS")
        arquivo_modelo = st.file_uploader(
            "Faça upload do modelo Word (.docx)",
            type=['docx'], 
            help="Modelo de OS com placeholders"
        )

    if arquivo_funcionarios and arquivo_modelo:
        try:
            df_funcionarios = pd.read_excel(arquivo_funcionarios)
            st.success(f"✅ Planilha carregada: **{len(df_funcionarios)} funcionários**")

            # Mostrar preview
            with st.expander("👀 Prévia da planilha"):
                st.dataframe(df_funcionarios.head())

            creditos_necessarios = len(df_funcionarios)

            if creditos_necessarios > user.get('credits', 0):
                st.error(f"❌ Créditos insuficientes! Necessário: **{creditos_necessarios}** | Disponível: **{user.get('credits', 0)}**")
            else:
                st.info(f"💰 Serão utilizados **{creditos_necessarios} créditos** para gerar {creditos_necessarios} OS")

                if st.button("🚀 GERAR ORDENS DE SERVIÇO", type="primary", use_container_width=True):
                    with st.spinner("Gerando ordens de serviço..."):
                        # Usar créditos
                        username = user['username']
                        st.session_state.users_db[username]['credits'] -= creditos_necessarios
                        st.session_state.user['credits'] -= creditos_necessarios

                        # Gerar documentos
                        documentos_gerados = processar_os_lote(df_funcionarios, arquivo_modelo)

                        if documentos_gerados:
                            # Criar ZIP
                            zip_data = criar_zip_documentos(documentos_gerados)

                            st.success(f"✅ {len(documentos_gerados)} ordens de serviço geradas com sucesso!")
                            st.balloons()

                            # Download
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            filename = f"OS_Lote_{timestamp}.zip"

                            st.download_button(
                                "📥 DOWNLOAD DAS OS GERADAS",
                                data=zip_data,
                                file_name=filename,
                                mime="application/zip",
                                type="primary",
                                use_container_width=True
                            )
                        else:
                            st.error("Erro ao gerar documentos")

        except Exception as e:
            st.error(f"Erro ao processar arquivos: {str(e)}")

    if st.button("← Voltar ao Dashboard"):
        st.session_state.page = 'dashboard'
        st.rerun()

# Aplicação principal
def main():
    init_session()

    page = st.session_state.get('page', 'login')

    if not st.session_state.logged_in:
        if page == 'forgot_password':
            show_login_page()  # Tab esqueci senha já incluída
        else:
            show_login_page()
    else:
        if page == 'dashboard':
            show_dashboard()
        elif page == 'generator':
            show_generator_page()
        else:
            show_dashboard()

if __name__ == "__main__":
    main()
