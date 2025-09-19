import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile
from io import BytesIO
import time
import re
from datetime import datetime
import os

# Configuração da página
st.set_page_config(
    page_title="Gerador de OS Profissional",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# CSS personalizado
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 2rem;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 2rem;
    }
    .info-card {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 4px solid #667eea;
        margin: 1rem 0;
    }
    .success-card {
        background: #d4edda;
        color: #155724;
        padding: 1.5rem;
        border-radius: 10px;
        border: 1px solid #c3e6cb;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

def gerar_os_do_funcionario(funcionario, modelo_docx):
    """Gera OS individual para um funcionário"""
    doc = Document(modelo_docx)

    # Dicionário de substituições
    substituicoes = {
        '{NOME_FUNCIONARIO}': str(funcionario.get('NOME', '')),
        '{FUNCAO}': str(funcionario.get('FUNCAO', '')),
        '{SETOR}': str(funcionario.get('SETOR', '')),
        '{DATA_AVALIACAO}': str(funcionario.get('DATA_AVALIACAO', '')),
        '{AGENTE_DE_RISCO}': str(funcionario.get('AGENTE_DE_RISCO', '')),
        '{FONTE_GERADORA}': str(funcionario.get('FONTE_GERADORA', '')),
        '{INTENSIDADE_CONCENTRACAO}': str(funcionario.get('INTENSIDADE_CONCENTRACAO', '')),
        '{UNIDADE_DE_MEDIDA}': str(funcionario.get('UNIDADE_DE_MEDIDA', '')),
        '{LIMITE_DE_TOLERANCIA}': str(funcionario.get('LIMITE_DE_TOLERANCIA', '')),
        '{TECNICA_UTILIZADA}': str(funcionario.get('TECNICA_UTILIZADA', '')),
        '{EPC_EXISTENTE}': str(funcionario.get('EPC_EXISTENTE', '')),
        '{EPC_RECOMENDADO}': str(funcionario.get('EPC_RECOMENDADO', '')),
        '{EPI_EXISTENTE}': str(funcionario.get('EPI_EXISTENTE', '')),
        '{EPI_RECOMENDADO}': str(funcionario.get('EPI_RECOMENDADO', '')),
        '{OBSERVACOES}': str(funcionario.get('OBSERVACOES', ''))
    }

    # Substituir nos parágrafos
    for paragraph in doc.paragraphs:
        for placeholder, value in substituicoes.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Substituir nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in substituicoes.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    return doc

def processar_os_lote(df_funcionarios, modelo_docx):
    """Processa lote de funcionários e gera todas as OS"""
    documentos_gerados = []

    progress_bar = st.progress(0)
    status_text = st.empty()

    for index, funcionario in df_funcionarios.iterrows():
        try:
            status_text.text(f"Gerando OS para: {funcionario.get('NOME', 'Funcionário')} ({index+1}/{len(df_funcionarios)})")

            doc_gerado = gerar_os_do_funcionario(funcionario, modelo_docx)

            # Gerar nome do arquivo
            nome_funcionario = str(funcionario.get('NOME', f'Funcionario_{index}')).strip()
            nome_funcionario = re.sub(r'[^\w\s-]', '', nome_funcionario).strip()
            nome_funcionario = re.sub(r'[-\s]+', '_', nome_funcionario)

            if not nome_funcionario:
                nome_funcionario = f'Funcionario_{index}'

            nome_arquivo = f"OS_{nome_funcionario}.docx"

            # Salvar documento em bytes
            doc_bytes = BytesIO()
            doc_gerado.save(doc_bytes)
            doc_bytes.seek(0)

            documentos_gerados.append({
                'nome': nome_arquivo,
                'conteudo': doc_bytes.getvalue()
            })

            progress_bar.progress((index + 1) / len(df_funcionarios))

        except Exception as e:
            st.error(f"Erro ao gerar OS para {funcionario.get('NOME', 'funcionário')}: {str(e)}")
            continue

    status_text.text("✅ Processamento concluído!")
    return documentos_gerados

def criar_zip_documentos(documentos):
    """Cria arquivo ZIP com todos os documentos"""
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for doc in documentos:
            zip_file.writestr(doc['nome'], doc['conteudo'])

    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def main():
    """Aplicação principal"""

    # Cabeçalho
    st.markdown("""
    <div class="main-header">
        <h1>🏗️ Gerador de OS Profissional</h1>
        <p>Sistema automatizado para geração de Ordens de Serviço</p>
        <small>⚡ Versão simples - Upload, processe e baixe suas OS!</small>
    </div>
    """, unsafe_allow_html=True)

    # Instruções
    with st.expander("📋 Como usar o sistema", expanded=False):
        st.markdown("""
        ### 🎯 Passo a passo:
        1. **📊 Upload da Planilha:** Faça upload do arquivo Excel (.xlsx) com os dados dos funcionários
        2. **📄 Upload do Modelo:** Faça upload do modelo de OS em Word (.docx)
        3. **🚀 Gerar OS:** Clique no botão para processar automaticamente
        4. **📥 Download:** Baixe o arquivo ZIP com todas as OS geradas

        ### 📝 Campos obrigatórios na planilha:
        - **NOME:** Nome do funcionário
        - **FUNCAO:** Função exercida
        - **SETOR:** Setor de trabalho
        - **DATA_AVALIACAO:** Data da avaliação
        - **AGENTE_DE_RISCO:** Tipo de risco identificado
        - **FONTE_GERADORA:** Fonte do risco
        - **INTENSIDADE_CONCENTRACAO:** Nível do risco
        - **UNIDADE_DE_MEDIDA:** Unidade de medida
        - **LIMITE_DE_TOLERANCIA:** Limite permitido
        - **TECNICA_UTILIZADA:** Técnica de avaliação
        - **EPC_EXISTENTE:** EPC atual
        - **EPC_RECOMENDADO:** EPC recomendado
        - **EPI_EXISTENTE:** EPI atual
        - **EPI_RECOMENDADO:** EPI recomendado
        - **OBSERVACOES:** Observações adicionais
        """)

    # Interface principal
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 📊 Planilha de Funcionários")
        arquivo_funcionarios = st.file_uploader(
            "Selecione a planilha Excel (.xlsx)",
            type=['xlsx'],
            help="Planilha com dados dos funcionários para geração das OS"
        )

        if arquivo_funcionarios:
            try:
                df_funcionarios = pd.read_excel(arquivo_funcionarios)

                st.markdown(f"""
                <div class="success-card">
                    ✅ <strong>Planilha carregada com sucesso!</strong><br>
                    📊 <strong>{len(df_funcionarios)} funcionários</strong> encontrados<br>
                    📋 <strong>{len(df_funcionarios.columns)} colunas</strong> detectadas
                </div>
                """, unsafe_allow_html=True)

                # Mostrar prévia
                with st.expander("👀 Prévia dos dados (5 primeiras linhas)"):
                    st.dataframe(df_funcionarios.head())

                # Verificar colunas obrigatórias
                colunas_obrigatorias = [
                    'NOME', 'FUNCAO', 'SETOR', 'DATA_AVALIACAO', 'AGENTE_DE_RISCO'
                ]

                colunas_faltando = [col for col in colunas_obrigatorias if col not in df_funcionarios.columns]

                if colunas_faltando:
                    st.error(f"❌ Colunas obrigatórias faltando: {', '.join(colunas_faltando)}")
                else:
                    st.success("✅ Todas as colunas obrigatórias encontradas!")

            except Exception as e:
                st.error(f"❌ Erro ao ler planilha: {str(e)}")
                arquivo_funcionarios = None

    with col2:
        st.markdown("### 📄 Modelo de OS")
        arquivo_modelo = st.file_uploader(
            "Selecione o modelo Word (.docx)", 
            type=['docx'],
            help="Modelo de OS com placeholders para substituição automática"
        )

        if arquivo_modelo:
            st.markdown("""
            <div class="success-card">
                ✅ <strong>Modelo carregado com sucesso!</strong><br>
                📄 Pronto para processamento
            </div>
            """, unsafe_allow_html=True)

            with st.expander("ℹ️ Sobre os placeholders"):
                st.markdown("""
                ### 🔧 Placeholders disponíveis:
                O sistema substituirá automaticamente estes códigos no seu modelo:

                - `{NOME_FUNCIONARIO}` → Nome do funcionário
                - `{FUNCAO}` → Função
                - `{SETOR}` → Setor
                - `{DATA_AVALIACAO}` → Data da avaliação
                - `{AGENTE_DE_RISCO}` → Agente de risco
                - `{FONTE_GERADORA}` → Fonte geradora
                - `{INTENSIDADE_CONCENTRACAO}` → Intensidade/concentração
                - `{UNIDADE_DE_MEDIDA}` → Unidade de medida
                - `{LIMITE_DE_TOLERANCIA}` → Limite de tolerância
                - `{TECNICA_UTILIZADA}` → Técnica utilizada
                - `{EPC_EXISTENTE}` → EPC existente
                - `{EPC_RECOMENDADO}` → EPC recomendado
                - `{EPI_EXISTENTE}` → EPI existente
                - `{EPI_RECOMENDADO}` → EPI recomendado
                - `{OBSERVACOES}` → Observações
                """)

    # Processamento
    if arquivo_funcionarios and arquivo_modelo:
        st.markdown("---")

        try:
            df_funcionarios = pd.read_excel(arquivo_funcionarios)

            # Verificar se tem dados
            if len(df_funcionarios) == 0:
                st.error("❌ A planilha está vazia!")
                return

            # Informações finais
            st.markdown(f"""
            <div class="info-card">
                🎯 <strong>Pronto para gerar:</strong><br>
                📊 <strong>{len(df_funcionarios)} Ordens de Serviço</strong><br>
                🕒 <strong>Tempo estimado:</strong> {len(df_funcionarios) * 2} segundos<br>
                📁 <strong>Formato de saída:</strong> ZIP com arquivos .docx individuais
            </div>
            """, unsafe_allow_html=True)

            # Botão de geração
            if st.button("🚀 GERAR ORDENS DE SERVIÇO", type="primary", use_container_width=True):

                with st.spinner("Processando ordens de serviço..."):
                    inicio = time.time()

                    # Processar documentos
                    documentos_gerados = processar_os_lote(df_funcionarios, arquivo_modelo)

                    if documentos_gerados:
                        # Criar ZIP
                        zip_data = criar_zip_documentos(documentos_gerados)

                        fim = time.time()
                        tempo_total = fim - inicio

                        # Nome do arquivo de saída
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        nome_zip = f"OS_Lote_{timestamp}.zip"

                        # Estatísticas
                        tamanho_zip = len(zip_data) / 1024  # KB

                        st.markdown(f"""
                        <div class="success-card">
                            🎉 <strong>Processamento concluído com sucesso!</strong><br><br>
                            📊 <strong>Estatísticas:</strong><br>
                            ✅ {len(documentos_gerados)} OS geradas<br>
                            🕒 Tempo total: {tempo_total:.1f} segundos<br>
                            📁 Tamanho do arquivo: {tamanho_zip:.1f} KB<br>
                            📄 Arquivo: {nome_zip}
                        </div>
                        """, unsafe_allow_html=True)

                        # Balões de comemoração
                        st.balloons()

                        # Botão de download
                        st.download_button(
                            label="📥 DOWNLOAD DAS OS GERADAS",
                            data=zip_data,
                            file_name=nome_zip,
                            mime="application/zip",
                            type="primary",
                            use_container_width=True
                        )

                        # Informações adicionais
                        with st.expander(f"📋 Detalhes dos {len(documentos_gerados)} arquivos gerados"):
                            for i, doc in enumerate(documentos_gerados, 1):
                                st.write(f"{i}. {doc['nome']}")

                    else:
                        st.error("❌ Nenhum documento foi gerado. Verifique os dados da planilha.")

        except Exception as e:
            st.error(f"❌ Erro durante o processamento: {str(e)}")

    # Rodapé
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 1rem;">
        <small>
        🏗️ <strong>Gerador de OS Profissional</strong> | 
        Desenvolvido para automatizar a geração de Ordens de Serviço | 
        ⚡ Processamento rápido e eficiente
        </small>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
