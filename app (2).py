# 📄 Gerador de Ordens de Serviço (OS)

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
import zipfile
from io import BytesIO
import time
import re
import datetime

# --- Configuração da Página ---
st.set_page_config(
    page_title="Gerador de Ordens de Serviço (OS)",
    page_icon="📄",
    layout="wide",
)

# --- DEFINIÇÃO DE CONSTANTES GLOBAIS ---
UNIDADES_DE_MEDIDA = ["dB(A)", "m/s²", "ppm", "mg/m³", "%", "°C", "lx", "cal/cm²", "µT", "kV/m", "W/m²", "f/cm³", "Não aplicável"]

# --- AGENTES DE RISCO ORGANIZADOS POR CATEGORIA ---
RISCOS_FISICO = sorted([
    "Ambiente Artificialmente Frio",
    "Exposição ao Ruído",
    "Exposição à Radiações Ionizantes",
    "Exposição à Radiações Não-ionizantes",
    "Exposição à Temperatura Ambiente Baixa",
    "Exposição à Temperatura Ambiente Elevada",
    "Pressão Atmosférica Anormal (condições hiperbáricas)",
    "Umidade",
    "Vibração de Corpo Inteiro (AREN)",
    "Vibração de Corpo Inteiro (VDVR)",
    "Vibrações Localizadas (mão/braço)"
])

RISCOS_QUIMICO = sorted([
    "Exposição a Produto Químico",
    "Poeiras",
    "Fumos",
    "Névoas", 
    "Neblinas",
    "Gases",
    "Vapores",
    "Produtos Químicos em Geral"
])

RISCOS_BIOLOGICO = sorted([
    "Água e/ou alimentos contaminados",
    "Contaminação pelo Corona Vírus",
    "Contato com Fluido Orgânico (sangue, hemoderivados, secreções, excreções)",
    "Contato com Pessoas Doentes e/ou Material Infectocontagiante",
    "Exposição à Agentes Microbiológicos (fungos, bactérias, vírus, protozoários, parasitas)",
    "Vírus",
    "Bactérias",
    "Protozoários",
    "Fungos",
    "Parasitas",
    "Bacilos"
])

RISCOS_ERGONOMICO = sorted([
    "Assento inadequado",
    "Assédio de qualquer natureza no trabalho",
    "Cadência do trabalho imposta por um equipamento",
    "Compressão de partes do corpo por superfícies rígidas ou com quinas vivas",
    "Conflitos hierárquicos no trabalho",
    "Desequilíbrio entre tempo de trabalho e tempo de repouso",
    "Dificuldades para cumprir ordens e determinações da chefia relacionadas ao trabalho",
    "Elevação frequente de membros superiores",
    "Encosto do assento inadequado ou ausente",
    "Equipamentos e/ou máquinas sem meios de regulagem de ajustes ou sem condições de uso",
    "Equipamentos/mobiliário não adaptados à antropometria do trabalhador",
    "Esforço físico intenso",
    "Exigência de concentração, atenção e memória",
    "Exposição à vibração de corpo inteiro",
    "Exposição à vibrações localizadas (mão, braço)",
    "Falta de autonomia para a realização de tarefas no trabalho",
    "Flexões da coluna vertebral frequentes",
    "Frequente ação de empurrar/puxar cargas ou volumes",
    "Frequente deslocamento à pé durante à jornada de trabalho",
    "Frequente execução de movimentos repetitivos",
    "Iluminação inadequada",
    "Insatisfação no trabalho",
    "Insuficiência de capacitação para à execução da tarefa",
    "Levantamento e transporte manual de cargas ou volumes",
    "Manuseio de ferramentas e/ou objetos pesados por longos períodos",
    "Manuseio ou movimentação de cargas e volumes sem pega ou com \"pega pobre\"",
    "Mobiliário ou equipamento sem espaço para movimentação de segmentos corporais",
    "Mobiliário sem meios de regulagem de ajustes",
    "Monotonia",
    "Necessidade de alcançar objetos, documentos, controles, etc, além das zonas de alcance ideais",
    "Necessidade de manter ritmos intensos de trabalho",
    "Piso escorregadio ou irregular",
    "Posto de trabalho improvisado/inadequado",
    "Posto de trabalho não planejado/adaptado para à posição sentada",
    "Postura em pé por longos períodos",
    "Postura sentada por longos períodos",
    "Posturas incômodas/pouco confortáveis por longos períodos",
    "Pressão sonora fora dos parâmetros de conforto",
    "Problemas de relacionamento no trabalho",
    "Realização de múltiplas tarefas com alta demanda mental/cognitiva",
    "Reflexos que causem desconforto ou prejudiquem à visão",
    "Situações de estresse no local de trabalho",
    "Situações de sobrecarga de trabalho mental",
    "Temperatura efetiva fora dos parâmetros de conforto",
    "Trabalho com necessidade de variação de turnos",
    "Trabalho com utilização rigorosa de metas de produção",
    "Trabalho em condições de difícil comunicação",
    "Trabalho intensivo com teclado ou outros dispositivos de entrada de dados",
    "Trabalho noturno",
    "Trabalho realizado sem pausas pré-definidas para descanso",
    "Trabalho remunerado por produção",
    "Umidade do ar fora dos parâmetros de conforto",
    "Uso frequente de alavancas",
    "Uso frequente de escadas",
    "Uso frequente de força, pressão, preensão, flexão, extensão ou torção dos segmentos corporais",
    "Uso frequente de pedais",
    "Velocidade do ar fora dos parâmetros de conforto"
])

RISCOS_ACIDENTE = sorted([
    "Absorção (por contato) de substância cáustica, tóxica ou nociva",
    "Afogamento, imersão, engolfamento",
    "Aprisionamento em, sob ou entre",
    "Aprisionamento em, sob ou entre desabamento ou desmoronamento de edificação, estrutura, barreira, etc",
    "Aprisionamento em, sob ou entre dois ou mais objetos em movimento (sem encaixe)",
    "Aprisionamento em, sob ou entre objetos em movimento convergente",
    "Aprisionamento em, sob ou entre um objeto parado e outro em movimento",
    "Arestas cortantes, superfícies com rebarbas, farpas ou elementos de fixação expostos",
    "Ataque de ser vivo (inclusive humano)",
    "Ataque de ser vivo com peçonha",
    "Ataque de ser vivo com transmissão de doença",
    "Ataque de ser vivo por mordedura, picada, chifrada, coice, etc",
    "Atrito ou abrasão",
    "Atrito ou abrasão por corpo estranho no olho",
    "Atrito ou abrasão por encostar em objeto",
    "Atrito ou abrasão por manusear objeto",
    "Atropelamento",
    "Batida contra objeto parado ou em movimento",
    "Carga Suspensa",
    "Colisão entre veículos e/ou equipamentos autopropelidos",
    "Condições climáticas adversas (sol, chuva, vento, etc)",
    "Contato com objeto ou substância a temperatura muito alta",
    "Contato com objeto ou substância a temperatura muito baixa",
    "Contato com objeto ou substância em movimento",
    "Desabamento/Desmoronamento de edificação, estrutura e/ou materiais diversos",
    "Elementos Móveis e/ou Rotativos",
    "Emergências na circunvizinhança",
    "Equipamento pressurizado hidráulico ou pressurizado",
    "Exposição à Energia Elétrica",
    "Ferramentas elétricas",
    "Ferramentas manuais",
    "Gases/vapores/poeiras (tóxicos ou não tóxicos)",
    "Gases/vapores/poeiras inflamáveis",
    "Impacto de pessoa contra objeto em movimento",
    "Impacto de pessoa contra objeto parado",
    "Impacto sofrido por pessoa",
    "Impacto sofrido por pessoa, de objeto em movimento",
    "Impacto sofrido por pessoa, de objeto projetado",
    "Impacto sofrido por pessoa, de objeto que cai",
    "Incêndio/Explosão",
    "Ingestão de substância cáustica, tóxica ou nociva",
    "Inalação de substância tóxica/nociva",
    "Inalação, ingestão e/ou absorção",
    "Objetos cortantes/perfurocortantes",
    "Pessoas não autorizadas e/ou visitantes no local de trabalho",
    "Portas, escotilhas, tampas, \"bocas de visita\", flanges",
    "Projeção de Partículas sólidas e/ou líquidas",
    "Queda de pessoa com diferença de nível > 2m",
    "Queda de pessoa com diferença de nível ≤ 2m",
    "Queda de pessoa com diferença de nível de andaime, passarela, plataforma, etc",
    "Queda de pessoa com diferença de nível de escada (móvel ou fixa)",
    "Queda de pessoa com diferença de nível de material empilhado",
    "Queda de pessoa com diferença de nível de veículo",
    "Queda de pessoa com diferença de nível em poço, escavação, abertura no piso, etc",
    "Queda de pessoa em mesmo nível",
    "Reação do corpo a seus movimentos (escorregão sem queda, etc)",
    "Soterramento",
    "Substâncias tóxicas e/ou inflamáveis",
    "Superfícies, substâncias e/ou objetos aquecidos",
    "Superfícies, substâncias e/ou objetos em baixa temperatura",
    "Tombamento de máquina/equipamento",
    "Tombamento, quebra e/ou ruptura de estrutura (fixa ou móvel)",
    "Trabalho à céu aberto",
    "Trabalho com máquinas e/ou equipamentos",
    "Trabalho com máquinas portáteis rotativas",
    "Trabalho em espaços confinados",
    "Vidro (recipientes, portas, bancadas, janelas, objetos diversos)"
])

# Dicionário para mapear categorias aos riscos
AGENTES_POR_CATEGORIA = {
    'fisico': RISCOS_FISICO,
    'quimico': RISCOS_QUIMICO,
    'biologico': RISCOS_BIOLOGICO,
    'ergonomico': RISCOS_ERGONOMICO,
    'acidente': RISCOS_ACIDENTE
}

CATEGORIAS_RISCO = {
    'fisico': '🔥 Físicos',
    'quimico': '⚗️ Químicos', 
    'biologico': '🦠 Biológicos',
    'ergonomico': '🏃 Ergonômicos',
    'acidente': '⚠️ Acidentes'
}

# --- CSS PERSONALIZADO ---
st.markdown("""
<style>
    .main {
        padding-top: 0rem;
    }
    
    .stApp > header {
        background-color: transparent;
    }
    
    .block-container {
        padding-top: 2rem;
        padding-bottom: 0rem;
        padding-left: 1rem;
        padding-right: 1rem;
    }
    
    h1 {
        color: #1f2937;
        text-align: center;
        font-size: 2.5rem;
        margin-bottom: 1rem;
    }
    
    .upload-box {
        background-color: #f8fafc;
        border: 2px dashed #3b82f6;
        border-radius: 0.5rem;
        padding: 2rem;
        text-align: center;
        margin: 1rem 0;
    }
    
    .success-box {
        background-color: #f0fdf4;
        border: 1px solid #bbf7d0;
        border-radius: 0.5rem;
        padding: 1rem;
        margin: 1rem 0;
    }
    
    .warning-box {
        background-color: #fefce8;
        border: 1px solid #fde047;
        border-radius: 0.5rem;
        padding: 1rem;
        margin: 1rem 0;
    }
    
    .info-box {
        background-color: #eff6ff;
        border: 1px solid #93c5fd;
        border-radius: 0.5rem;
        padding: 1rem;
        margin: 1rem 0;
    }
    
    .metric-container {
        background-color: white;
        border: 1px solid #e5e7eb;
        border-radius: 0.5rem;
        padding: 1rem;
        text-align: center;
        margin: 0.5rem 0;
    }
    
    .metric-container h3 {
        color: #1f2937;
        font-size: 2rem;
        margin-bottom: 0.5rem;
    }
    
    .metric-container p {
        color: #6b7280;
        font-size: 0.9rem;
        margin: 0;
    }
    
    .stButton > button {
        width: 100%;
        border-radius: 0.5rem;
        border: none;
        padding: 0.75rem 1rem;
        font-weight: 500;
        transition: all 0.2s;
    }
    
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15);
    }
    
    .stSelectbox > div > div {
        border-radius: 0.5rem;
    }
    
    .stTextInput > div > div {
        border-radius: 0.5rem;
    }
    
    .stTextArea > div > div {
        border-radius: 0.5rem;
    }
    
    .stMultiSelect > div > div {
        border-radius: 0.5rem;
    }
    
    .streamlit-expanderHeader {
        background-color: #f8fafc;
        border-radius: 0.5rem;
        border: 1px solid #e5e7eb;
    }
    
    .streamlit-expanderContent {
        border: 1px solid #e5e7eb;
        border-top: none;
        border-radius: 0 0 0.5rem 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

# --- FUNÇÕES AUXILIARES ---
def create_sample_data():
    """Cria dados de exemplo para demonstração"""
    sample_data = {
        'Nome': ['JOÃO SILVA SANTOS', 'MARIA OLIVEIRA COSTA', 'PEDRO ALVES FERREIRA'],
        'Setor': ['PRODUCAO DE LA DE ACO', 'ADMINISTRACAO DE RH', 'MANUTENCAO QUIMICA'],
        'Função': ['OPERADOR PRODUCAO I', 'ANALISTA ADM PESSOAL PL', 'MECANICO MANUT II'],
        'Data de Admissão': ['15/03/2020', '22/08/2019', '10/01/2021'],
        'Empresa': ['SUA EMPRESA', 'SUA EMPRESA', 'SUA EMPRESA'],
        'Unidade': ['Matriz', 'Matriz', 'Matriz'],
        'Descrição de Atividades': [
            'Operar equipamentos de produção nível I, controlar parâmetros operacionais, realizar inspeções visuais e registrar dados de produção.',
            'Executar atividades de administração de pessoal, controlar documentos trabalhistas, elaborar relatórios e dar suporte às equipes.',
            'Executar manutenção preventiva e corretiva em equipamentos, diagnosticar falhas, trocar componentes e registrar intervenções.'
        ]
    }
    return pd.DataFrame(sample_data)

def validate_excel_structure(df):
    """Valida se a planilha tem a estrutura necessária"""
    required_columns = ['Nome', 'Setor', 'Função', 'Data de Admissão', 'Empresa', 'Unidade', 'Descrição de Atividades']
    missing_columns = [col for col in required_columns if col not in df.columns]
    
    if missing_columns:
        return False, f"Colunas obrigatórias faltando: {', '.join(missing_columns)}"
    
    if df.empty:
        return False, "A planilha está vazia"
    
    return True, "Estrutura válida"

def gerar_documento_os(dados_funcionario, agentes_risco, epis, medidas_preventivas, observacoes, template_doc=None):
    """
    Gera a Ordem de Serviço com base nos dados fornecidos
    """
    try:
        # Usar template personalizado se fornecido, senão criar um padrão
        if template_doc:
            doc = template_doc
        else:
            doc = Document()
        
        # Se for um template novo, adicionar estrutura básica
        if not template_doc:
            # Título
            titulo = doc.add_heading('ORDEM DE SERVIÇO', 0)
            titulo.alignment = 1  # Centralizado
            
            # Subtítulo
            subtitulo = doc.add_paragraph('Informações sobre Condições de Segurança e Saúde no Trabalho - NR-01')
            subtitulo.alignment = 1
            
            doc.add_paragraph()  # Espaço
        
        # Informações do Funcionário
        info_func = doc.add_paragraph()
        info_func.add_run(f"Empresa: {dados_funcionario.get('Empresa', '')}\t\t")
        info_func.add_run(f"Unidade: {dados_funcionario.get('Unidade', '')}")
        
        info_func2 = doc.add_paragraph()
        info_func2.add_run(f"Nome do Funcionário: {dados_funcionario.get('Nome', '')}")
        
        info_func3 = doc.add_paragraph()
        info_func3.add_run(f"Data de Admissão: {dados_funcionario.get('Data de Admissão', '')}\t\t")
        
        info_func4 = doc.add_paragraph()
        info_func4.add_run(f"Setor de Trabalho: {dados_funcionario.get('Setor', '')}\t\t")
        info_func4.add_run(f"Função: {dados_funcionario.get('Função', '')}")
        
        doc.add_paragraph()  # Espaço
        
        # Tarefas da Função
        doc.add_heading('TAREFAS DA FUNÇÃO', level=1)
        doc.add_paragraph(dados_funcionario.get('Descrição de Atividades', 'Atividades relacionadas à função exercida.'))
        
        # Agentes de Riscos Ocupacionais
        if agentes_risco:
            doc.add_heading('AGENTES DE RISCOS OCUPACIONAIS', level=1)
            
            for categoria, riscos in agentes_risco.items():
                if riscos:  # Se há riscos nesta categoria
                    categoria_titulo = categoria.replace('_', ' ').title()
                    doc.add_heading(f'Riscos {categoria_titulo}', level=2)
                    
                    for risco in riscos:
                        risco_para = doc.add_paragraph()
                        risco_para.add_run(f"• {risco['agente']}")
                        if risco.get('intensidade'):
                            risco_para.add_run(f": {risco['intensidade']}")
                        if risco.get('unidade'):
                            risco_para.add_run(f" {risco['unidade']}")
        
        # EPIs Obrigatórios
        if epis:
            doc.add_heading('EQUIPAMENTOS DE PROTEÇÃO INDIVIDUAL (EPIs)', level=1)
            for epi in epis:
                doc.add_paragraph(f"• {epi}", style='List Bullet')
        
        # Medidas Preventivas
        if medidas_preventivas:
            doc.add_heading('MEDIDAS PREVENTIVAS E DE CONTROLE', level=1)
            for medida in medidas_preventivas:
                doc.add_paragraph(f"• {medida}", style='List Bullet')
        
        # Procedimentos de Emergência
        doc.add_heading('PROCEDIMENTOS EM SITUAÇÕES DE EMERGÊNCIA', level=1)
        emergencia_texto = """
• Comunique imediatamente o acidente à chefia imediata ou responsável pela área;
• Preserve as condições do local de acidente até a comunicação com a autoridade competente;
• Procure atendimento médico no ambulatório da empresa ou serviço médico de emergência;
• Siga as orientações do Plano de Emergência da empresa;
• Registre a ocorrência conforme procedimentos estabelecidos.
        """
        doc.add_paragraph(emergencia_texto.strip())
        
        # Grave e Iminente Risco
        doc.add_heading('ORIENTAÇÕES SOBRE GRAVE E IMINENTE RISCO', level=1)
        gir_texto = """
• Sempre que constatar condição de grave e iminente risco, interrompa imediatamente as atividades;
• Comunique de forma urgente ao seu superior hierárquico;
• Aguarde as providências necessárias e autorização para retorno;
• É direito do trabalhador recusar-se a trabalhar em condições de risco grave e iminente.
        """
        doc.add_paragraph(gir_texto.strip())
        
        # Observações Adicionais
        if observacoes:
            doc.add_heading('OBSERVAÇÕES COMPLEMENTARES', level=1)
            doc.add_paragraph(observacoes)
        
        # Nota Legal
        doc.add_paragraph()
        nota_legal = doc.add_paragraph()
        nota_legal.add_run("IMPORTANTE: ").bold = True
        nota_legal.add_run(
            "Conforme Art. 158 da CLT e NR-01, o descumprimento das disposições "
            "sobre segurança e saúde no trabalho sujeita o empregado às penalidades "
            "legais, inclusive demissão por justa causa."
        )
        
        # Assinaturas
        doc.add_paragraph()
        doc.add_paragraph("_" * 40 + "\t\t" + "_" * 40)
        doc.add_paragraph("Funcionário\t\t\t\t\tResponsável pela Área")
        doc.add_paragraph(f"Data: {datetime.date.today().strftime('%d/%m/%Y')}")
        
        return doc
        
    except Exception as e:
        st.error(f"Erro ao gerar documento: {str(e)}")
        return None

def main():
    # Título principal
    st.markdown("# 📄 Gerador de Ordens de Serviço (OS)")
    st.markdown("### Sistema para geração automatizada de OS conforme NR-01")
    
    # Sidebar com informações
    with st.sidebar:
        st.markdown("## 📋 Informações do Sistema")
        
        st.markdown(f"""
        **Funcionalidades:**
        - ✅ Upload de planilha Excel
        - ✅ **{sum(len(riscos) for riscos in AGENTES_POR_CATEGORIA.values())} riscos** ocupacionais
        - ✅ Geração individual ou em lote
        - ✅ Conformidade com NR-01
        - ✅ Download automático
        """)
        
        st.markdown("## 📊 Base de Riscos")
        for categoria, nome in CATEGORIAS_RISCO.items():
            qtd_riscos = len(AGENTES_POR_CATEGORIA[categoria])
            st.markdown(f"- {nome}: **{qtd_riscos}** opções")
        
        st.markdown("## 📁 Estrutura da Planilha")
        st.markdown("""
        **Colunas obrigatórias:**
        - Nome
        - Setor
        - Função
        - Data de Admissão
        - Empresa
        - Unidade
        - Descrição de Atividades
        """)
        
        # Botão para baixar planilha exemplo
        sample_df = create_sample_data()
        sample_buffer = BytesIO()
        sample_df.to_excel(sample_buffer, index=False)
        sample_buffer.seek(0)
        
        st.download_button(
            "📥 Baixar Planilha Exemplo",
            data=sample_buffer.getvalue(),
            file_name="modelo_funcionarios.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
    
    # Seção de Upload de Arquivos
    st.markdown("## 📤 Upload de Arquivos")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        <div class="upload-box">
            <h4>📊 Planilha de Funcionários</h4>
            <p>Faça upload da planilha Excel com os dados dos funcionários</p>
        </div>
        """, unsafe_allow_html=True)
        
        uploaded_excel = st.file_uploader(
            "Selecione a planilha Excel (.xlsx)",
            type=['xlsx'],
            help="A planilha deve conter todas as colunas obrigatórias listadas na barra lateral"
        )
    
    with col2:
        st.markdown("""
        <div class="upload-box">
            <h4>📄 Template de OS (Opcional)</h4>
            <p>Faça upload do seu template personalizado de Word</p>
        </div>
        """, unsafe_allow_html=True)
        
        uploaded_template = st.file_uploader(
            "Selecione o template Word (.docx) - Opcional",
            type=['docx'],
            help="Se não fornecido, será usado o template padrão do sistema"
        )
    
    # Processar planilha se foi carregada
    if uploaded_excel is not None:
        try:
            # Carregar e validar planilha
            df = pd.read_excel(uploaded_excel)
            is_valid, message = validate_excel_structure(df)
            
            if not is_valid:
                st.markdown(f"""
                <div class="warning-box">
                    <h4>⚠️ Erro na Validação da Planilha</h4>
                    <p>{message}</p>
                    <p>Por favor, corrija a planilha e faça upload novamente.</p>
                </div>
                """, unsafe_allow_html=True)
                return
            
            # Mostrar sucesso e estatísticas
            st.markdown(f"""
            <div class="success-box">
                <h4>✅ Planilha Carregada com Sucesso!</h4>
                <p>Foram encontrados <strong>{len(df)}</strong> funcionários na planilha.</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Exibir estatísticas
            st.markdown("### 📊 Estatísticas da Planilha")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.markdown(f"""
                <div class="metric-container">
                    <h3>{len(df)}</h3>
                    <p>👥 Funcionários</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="metric-container">
                    <h3>{df['Setor'].nunique()}</h3>
                    <p>🏢 Setores</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                st.markdown(f"""
                <div class="metric-container">
                    <h3>{df['Função'].nunique()}</h3>
                    <p>💼 Funções</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col4:
                st.markdown(f"""
                <div class="metric-container">
                    <h3>{df['Empresa'].nunique()}</h3>
                    <p>🏭 Empresas</p>
                </div>
                """, unsafe_allow_html=True)
            
            # Seleção de Funcionários
            st.markdown("## 👥 Seleção de Funcionários")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Filtro por setor
                setores_disponiveis = ['Todos os setores'] + sorted(df['Setor'].dropna().unique().tolist())
                setor_selecionado = st.selectbox(
                    "🏢 Filtrar por Setor:",
                    setores_disponiveis,
                    help="Selecione um setor específico ou mantenha 'Todos os setores'"
                )
                
                # Filtrar dataframe baseado no setor
                if setor_selecionado == 'Todos os setores':
                    df_filtrado = df
                else:
                    df_filtrado = df[df['Setor'] == setor_selecionado]
            
            with col2:
                # Modo de seleção
                modo_selecao = st.radio(
                    "📋 Modo de Seleção:",
                    ["Funcionário Individual", "Múltiplos Funcionários", "Todos do Setor Filtrado"],
                    help="Escolha como deseja selecionar os funcionários para gerar as OS"
                )
            
            # Lógica de seleção baseada no modo escolhido
            funcionarios_selecionados = []
            
            if modo_selecao == "Funcionário Individual":
                funcionario_individual = st.selectbox(
                    "👤 Selecione o funcionário:",
                    [''] + df_filtrado['Nome'].tolist(),
                    help="Escolha um funcionário específico"
                )
                if funcionario_individual:
                    funcionarios_selecionados = [funcionario_individual]
            
            elif modo_selecao == "Múltiplos Funcionários":
                funcionarios_selecionados = st.multiselect(
                    "👥 Selecione múltiplos funcionários:",
                    df_filtrado['Nome'].tolist(),
                    help="Escolha vários funcionários mantendo Ctrl pressionado"
                )
            
            else:  # Todos do setor
                funcionarios_selecionados = df_filtrado['Nome'].tolist()
                if funcionarios_selecionados:
                    st.info(f"📝 Serão geradas OS para todos os {len(funcionarios_selecionados)} funcionários do setor filtrado.")
            
            # Mostrar funcionários selecionados
            if funcionarios_selecionados:
                st.success(f"✅ {len(funcionarios_selecionados)} funcionário(s) selecionado(s) para geração de OS")
                
                # Configuração de Riscos e Medidas
                st.markdown("## ⚠️ Configuração de Riscos Ocupacionais")
                
                # Inicializar session state se necessário
                if 'agentes_risco' not in st.session_state:
                    st.session_state.agentes_risco = {categoria: [] for categoria in CATEGORIAS_RISCO.keys()}
                if 'epis_selecionados' not in st.session_state:
                    st.session_state.epis_selecionados = []
                if 'medidas_preventivas' not in st.session_state:
                    st.session_state.medidas_preventivas = []
                
                # Configuração de Agentes de Risco por categoria
                st.markdown("### 🔍 Agentes de Riscos por Categoria")
                
                for categoria_key, categoria_nome in CATEGORIAS_RISCO.items():
                    qtd_opcoes = len(AGENTES_POR_CATEGORIA[categoria_key])
                    with st.expander(f"{categoria_nome} ({qtd_opcoes} opções disponíveis)", expanded=False):
                        st.markdown(f"**Configurar {categoria_nome}**")
                        
                        col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
                        
                        with col1:
                            agente_selecionado = st.selectbox(
                                "Agente de Risco:",
                                ['Selecione...'] + AGENTES_POR_CATEGORIA[categoria_key],
                                key=f"agente_{categoria_key}"
                            )
                        
                        with col2:
                            intensidade = st.text_input(
                                "Intensidade/Valor:",
                                key=f"intensidade_{categoria_key}",
                                placeholder="Ex: 85"
                            )
                        
                        with col3:
                            unidade = st.selectbox(
                                "Unidade:",
                                UNIDADES_DE_MEDIDA,
                                key=f"unidade_{categoria_key}"
                            )
                        
                        with col4:
                            if st.button(f"➕ Adicionar", key=f"add_{categoria_key}"):
                                if agente_selecionado != 'Selecione...':
                                    novo_risco = {
                                        'agente': agente_selecionado,
                                        'intensidade': intensidade,
                                        'unidade': unidade
                                    }
                                    st.session_state.agentes_risco[categoria_key].append(novo_risco)
                                    st.success(f"✅ Risco {categoria_nome.lower()} adicionado!")
                                    st.rerun()
                        
                        # Mostrar riscos já adicionados
                        if st.session_state.agentes_risco[categoria_key]:
                            st.markdown("**Riscos configurados:**")
                            for idx, risco in enumerate(st.session_state.agentes_risco[categoria_key]):
                                col1, col2 = st.columns([5, 1])
                                with col1:
                                    risco_text = f"• {risco['agente']}"
                                    if risco['intensidade']:
                                        risco_text += f": {risco['intensidade']}"
                                    if risco['unidade'] and risco['unidade'] != 'Não aplicável':
                                        risco_text += f" {risco['unidade']}"
                                    st.write(risco_text)
                                with col2:
                                    if st.button("🗑️", key=f"remove_{categoria_key}_{idx}", help="Remover este risco"):
                                        st.session_state.agentes_risco[categoria_key].pop(idx)
                                        st.rerun()
                
                # Configuração de EPIs
                st.markdown("### 🥽 Equipamentos de Proteção Individual (EPIs)")
                
                col1, col2 = st.columns([4, 1])
                with col1:
                    novo_epi = st.text_input(
                        "Adicionar EPI:",
                        placeholder="Ex: Capacete de segurança, Óculos de proteção, etc."
                    )
                with col2:
                    if st.button("➕ Adicionar EPI"):
                        if novo_epi.strip():
                            st.session_state.epis_selecionados.append(novo_epi.strip())
                            st.success("✅ EPI adicionado!")
                            st.rerun()
                
                # Mostrar EPIs configurados
                if st.session_state.epis_selecionados:
                    st.markdown("**EPIs configurados:**")
                    for idx, epi in enumerate(st.session_state.epis_selecionados):
                        col1, col2 = st.columns([5, 1])
                        with col1:
                            st.write(f"• {epi}")
                        with col2:
                            if st.button("🗑️", key=f"remove_epi_{idx}", help="Remover este EPI"):
                                st.session_state.epis_selecionados.pop(idx)
                                st.rerun()
                
                # Configuração de Medidas Preventivas
                st.markdown("### 🛡️ Medidas Preventivas e de Controle")
                
                col1, col2 = st.columns([4, 1])
                with col1:
                    nova_medida = st.text_area(
                        "Adicionar Medida Preventiva:",
                        placeholder="Ex: Realizar pausas programadas durante a jornada de trabalho...",
                        height=80
                    )
                with col2:
                    st.write("")  # Espaço
                    st.write("")  # Espaço  
                    if st.button("➕ Adicionar Medida"):
                        if nova_medida.strip():
                            st.session_state.medidas_preventivas.append(nova_medida.strip())
                            st.success("✅ Medida preventiva adicionada!")
                            st.rerun()
                
                # Mostrar medidas configuradas
                if st.session_state.medidas_preventivas:
                    st.markdown("**Medidas preventivas configuradas:**")
                    for idx, medida in enumerate(st.session_state.medidas_preventivas):
                        col1, col2 = st.columns([5, 1])
                        with col1:
                            # Mostrar apenas os primeiros 100 caracteres
                            medida_resumida = medida[:100] + "..." if len(medida) > 100 else medida
                            st.write(f"• {medida_resumida}")
                        with col2:
                            if st.button("🗑️", key=f"remove_medida_{idx}", help="Remover esta medida"):
                                st.session_state.medidas_preventivas.pop(idx)
                                st.rerun()
                
                # Observações Complementares
                st.markdown("### 📝 Observações Complementares")
                observacoes = st.text_area(
                    "Informações adicionais para incluir nas OS:",
                    placeholder="Ex: Informações específicas do setor, procedimentos especiais, etc.",
                    height=100
                )
                
                # Botão para gerar OS
                st.markdown("## 🚀 Gerar Ordens de Serviço")
                
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button(
                        f"📄 GERAR {len(funcionarios_selecionados)} ORDEM(NS) DE SERVIÇO",
                        type="primary",
                        use_container_width=True
                    ):
                        # Carregar template se fornecido
                        template_doc = None
                        if uploaded_template:
                            try:
                                template_doc = Document(uploaded_template)
                            except Exception as e:
                                st.warning(f"⚠️ Erro ao carregar template personalizado: {str(e)}. Usando template padrão.")
                        
                        # Preparar dados para geração
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        documentos_gerados = []
                        
                        # Processar cada funcionário selecionado
                        for idx, nome_funcionario in enumerate(funcionarios_selecionados):
                            status_text.text(f"🔄 Gerando OS para: {nome_funcionario}")
                            
                            # Buscar dados do funcionário
                            dados_funcionario = df_filtrado[df_filtrado['Nome'] == nome_funcionario].iloc[0].to_dict()
                            
                            # Gerar documento
                            doc = gerar_documento_os(
                                dados_funcionario=dados_funcionario,
                                agentes_risco=st.session_state.agentes_risco,
                                epis=st.session_state.epis_selecionados,
                                medidas_preventivas=st.session_state.medidas_preventivas,
                                observacoes=observacoes,
                                template_doc=template_doc
                            )
                            
                            if doc:
                                # Salvar documento em buffer
                                buffer = BytesIO()
                                doc.save(buffer)
                                buffer.seek(0)
                                
                                documentos_gerados.append({
                                    'nome': nome_funcionario.replace(' ', '_').replace('/', '_'),
                                    'buffer': buffer
                                })
                            
                            # Atualizar progresso
                            progress_bar.progress((idx + 1) / len(funcionarios_selecionados))
                            time.sleep(0.1)  # Pequena pausa para visualização
                        
                        status_text.text("✅ Geração concluída!")
                        
                        # Disponibilizar downloads
                        if documentos_gerados:
                            if len(documentos_gerados) == 1:
                                # Download único
                                st.success(f"✅ Ordem de Serviço gerada com sucesso!")
                                st.download_button(
                                    label="📥 Download da Ordem de Serviço",
                                    data=documentos_gerados[0]['buffer'].getvalue(),
                                    file_name=f"OS_{documentos_gerados[0]['nome']}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                            else:
                                # Download em lote (ZIP)
                                st.success(f"✅ {len(documentos_gerados)} Ordens de Serviço geradas com sucesso!")
                                
                                # Criar arquivo ZIP
                                zip_buffer = BytesIO()
                                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                    for doc_info in documentos_gerados:
                                        zip_file.writestr(
                                            f"OS_{doc_info['nome']}.docx",
                                            doc_info['buffer'].getvalue()
                                        )
                                
                                zip_buffer.seek(0)
                                
                                st.download_button(
                                    label=f"📥 Download de {len(documentos_gerados)} Ordens de Serviço (ZIP)",
                                    data=zip_buffer.getvalue(),
                                    file_name=f"Lote_OS_{datetime.date.today().strftime('%d%m%Y')}.zip",
                                    mime="application/zip",
                                    use_container_width=True
                                )
                        else:
                            st.error("❌ Erro: Nenhum documento foi gerado. Verifique as configurações.")
        
        except Exception as e:
            st.markdown(f"""
            <div class="warning-box">
                <h4>❌ Erro no Processamento</h4>
                <p>Ocorreu um erro ao processar a planilha: <code>{str(e)}</code></p>
                <p>Verifique se o arquivo está correto e tente novamente.</p>
            </div>
            """, unsafe_allow_html=True)
    
    else:
        # Instruções iniciais quando não há planilha carregada
        st.markdown(f"""
        <div class="info-box">
            <h4>🎯 Como usar o sistema:</h4>
            <ol>
                <li><strong>📥 Baixe a planilha exemplo</strong> na barra lateral</li>
                <li><strong>✏️ Preencha</strong> com os dados dos seus funcionários</li>
                <li><strong>📤 Faça upload</strong> da planilha preenchida</li>
                <li><strong>⚙️ Configure</strong> os riscos ocupacionais e medidas de proteção</li>
                <li><strong>👥 Selecione</strong> os funcionários (individual, múltiplos ou todos)</li>
                <li><strong>📄 Gere</strong> as Ordens de Serviço conforme NR-01</li>
            </ol>
            
            <p><strong>🆕 Novidade:</strong> Agora com <strong>{sum(len(riscos) for riscos in AGENTES_POR_CATEGORIA.values())} opções de riscos</strong> ocupacionais organizados por categoria!</p>
            <p><strong>💡 Dica:</strong> Você também pode fazer upload de um template personalizado de Word para manter o padrão visual da sua empresa!</p>
        </div>
        """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()