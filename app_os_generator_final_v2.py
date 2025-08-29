import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
import zipfile
from io import BytesIO
import base64
import tempfile
import time

# --- Configuração da Página ---
st.set_page_config(
    page_title="Gerador de Ordens de Serviço (OS)",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- Inicialização do Session State ---
if 'descricoes' not in st.session_state:
    st.session_state.descricoes = {}
if 'documentos_gerados' not in st.session_state:
    st.session_state.documentos_gerados = []

# --- Modelo de OS Incorporado ---
MODELO_OS_TEMPLATE = """ORDEM DE SERVIÇO SOBRE SEGURANÇA E SAÚDE NO TRABALHO
NR01 item 1.4.1 c) item 1.4.4.1 b
N° da OS: Data de Elaboração: Última Revisão: Versão:
Pela presente Ordem de serviço, objetivamos informar os trabalhadores que executam suas atividades laborais
nesse setor, conforme estabelece a NR-01, sobre as condições de segurança e saúde às quais estão expostos, de
forma a padronizar comportamentos para prevenir acidentes e/ou doenças ocupacionais.
Empresa: [NOME EMPRESA] Unidade: [UNIDADE]
Nome do Funcionário: [NOME FUNCIONÁRIO] Data de Admissão: [DATA DE ADMISSÃO]
Setor de Trabalho: [SETOR] Função: [FUNÇÃO]
TAREFAS DA FUNÇÃO
[DESCRIÇÃO DE ATIVIDADES]
AGENTES DE RISCOS OCUPACIONAIS - NR01 item 1.4.1 b) I / item 1.4.4 a)
Físico: [RISCOS FÍSICOS]
Acidente: [RISCOS DE ACIDENTE]
Químico: [RISCOS QUÍMICOS]
Biológico: [RISCOS BIOLÓGICOS]
Ergonômicos: [RISCOS ERGONÔMICOS]
POSSÍVEIS DANOS À SAÚDE - NR01 item 1.4.1 b) I.
Físico: [POSSÍVEIS DANOS RISCOS FÍSICOS]
Acidente: [POSSÍVEIS DANOS RISCOS ACIDENTE]
Químico: [POSSÍVEIS DANOS RISCOS QUÍMICOS]
Biológico: [POSSÍVEIS DANOS RISCOS BIOLÓGICOS]
Ergonômicos: [POSSÍVEIS DANOS RISCOS ERGONÔMICOS]
MEIOS PARA O EMPREGADO PREVENIR E CONTROLAR OS RISCOS OCUPACIONAIS - NR01 item 1.4.4 b)
Barreira física, Protetor auricular silicone tipo plug; Insufladores / Exaustores / Ventilação natural / Pausas durante
a jornada / Hidratação disponível; Pausas programadas, detalhamento com AET do Trabalho, banco com anatomia
ergonômica Veículos industriais com dispositivos de segurança, sinalizador de ré, manutenção preventiva /
Operadores treinados e habilitados / Sinalização de atenção / Colete refletivo para DPA, Calçado de segurança /
Deslocamento com atenção / Proibição de caminhar abaixo de carga suspensa / Atenção durante a realização das
atividades; Barreira fisica, invólucros, Óculos de proteção para projeção de partículas.
MEDIDAS ADOTADAS PELA EMPRESA PARA REDUZIR OS RISCOS OCUPACIONAIS NR01 item 1.4.1 b) II /
item 1.4.4 c)
EPI: [EPIS]
Treinamento e Supervisão para execução das tarefas e uso dos EPI, em especial em relação aos trabalhos em
altura, com poeiras e solventes; Guarda-corpo de proteção periferias; Monitoramento do ambiente do trabalho
afim de corrigir condições inseguras encontradas, imediatamente; Fornecimento de cópia de ASO informando os
resultados dos exames médicos e dos exames complementares de diagnóstico aos quais os próprios
trabalhadores forem submetidos; Sinalização de Segurança no ambiente de trabalho; Fornecimento, Treinamento
e Exigência de uso de EPI.
INFORME DOS RESULTADOS DAS AVALIAÇÕES AMBIENTAIS NOS LOCAIS DE TRABALHO - NR01 item
1.4.1 b) IV.
[MEDIÇÕES]
PROCEDIMENTOS A SEREM ADOTADOS EM SITUAÇÃO DE ACIDENTES E EMERGÊNCIAS - NR01 item 1.4.4
d) / item 1.4.1 e)
Comunique imediatamente o acidente à chefia imediata ou na impossibilidade à pessoa que possa acessá-la;
Preserve as condições do local de acidente até a comunicação com a autoridade competente; Siga as orientações
correspondentes ao acidente e com as atribuições de sua função, indicados no "Plano de Respostas aos Possíveis
Cenários de Emergência", elaborado pela empresa".
ORIENTAÇÕES SOBRE CONSTATAÇÃO DE GRAVE E IMINENTE RISCO - NR01 item 1.4.4 e) / item 1.4.3 /
item 1.4.3.1
Sempre que constatar Grave e Iminente Risco à Vida e/ou Saúde, sua ou de outros, interrompa de imediato e com
segurança as atividades; Informe imediatamente ao seu superior hierárquico; Registre a constatação e as
medidas tomadas no "Registro de Condições de Grave e Iminente Risco", conforme procedimento padronizados
pela empresa; Aguarde as providências e liberação formal do cenário pela empresa.
Conforme Art. 158 da CLT e NR-01 item 1.4.2.1, o descumprimento imotivado das disposições legais e
regulamentares sobre segurança e saúde no trabalho, inclusive das ordens de serviço expedidas pelo
empregador, sujeita o empregado às penalidades legais, inclusive, demissão por justa causa.
X X X
SESMT Chefia Imediata Funcionário"""

# --- Unidades de Medida ---
UNIDADES_MEDIDA = [
    "dB Linear", "dB(C)", "dB(A)", "m/s²", "m/s1,75", "ppm", "mg/m³", "g/m³", 
    "f/cm³", "°C", "m/s", "%", "lx", "ufc/m³", "W/m²", "A/m", "mT", "µT", 
    "mA", "kV/m", "V/m", "J/m²", "mJ/cm²", "mSv", "mppdc", "UR(%)", "Lux"
]

# --- Funções de Lógica de Negócio ---

def normalizar_colunas(df):
    """Normaliza os nomes das colunas de um DataFrame."""
    if df is None:
        return None
    df.columns = (
        df.columns.str.lower()
        .str.strip()
        .str.replace(" ", "_")
        .str.replace("ç", "c").str.replace("ã", "a").str.replace("é", "e")
        .str.normalize("NFKD").str.encode("ascii", errors="ignore").str.decode("utf-8")
    )
    # Renomeia colunas específicas da planilha PGR para um padrão interno
    rename_map = {
        'perigo__(fator_de_risco/agente_nocivo/situacao_perigosa)': 'risco',
        'perigo_(fator_de_risco/agente_nocivo/situacao_perigosa)': 'risco',
        'possiveis_danos_ou_agravos_a_saude': 'possiveis_danos'
    }
    df = df.rename(columns=rename_map)

    if 'categoria' in df.columns:
        df['categoria'] = df['categoria'].str.normalize("NFKD").str.encode("ascii", errors="ignore").str.decode("utf-8").str.lower()
    return df

def mapear_e_renomear_colunas_funcionarios(df):
    """Tenta adivinhar e renomear colunas da planilha de funcionários."""
    mapeamento = {
        'nome_do_funcionario': ['nome', 'funcionario', 'colaborador', 'nome_completo'],
        'funcao': ['funcao', 'cargo', 'carga'],
        'data_de_admissao': ['data_admissao', 'data_de_admissao', 'admissao'],
        'setor': ['setor', 'area', 'departamento'],
        'matricula': ['matricula', 'registro', 'id'],
        'descricao_de_atividades': ['descricao_de_atividades', 'atividades', 'descricao_atividades', 'tarefas', 'funcoes'],
    }
    colunas_renomeadas = {}
    for nome_padrao, nomes_possiveis in mapeamento.items():
        for nome_possivel in nomes_possiveis:
            if nome_possivel in df.columns:
                colunas_renomeadas[nome_possivel] = nome_padrao
                break
    df = df.rename(columns=colunas_renomeadas)
    return df

@st.cache_data
def carregar_planilha(arquivo):
    """Carrega e processa uma planilha genérica."""
    if arquivo is None:
        return None
    try:
        df = pd.read_excel(arquivo)
        df = normalizar_colunas(df)
        return df
    except Exception as e:
        st.error(f"Erro ao ler o ficheiro Excel: {e}")
        return None

def criar_modelo_os_temporario():
    """Cria um arquivo temporário com o modelo de OS incorporado."""
    doc = Document()
    
    # Adiciona o conteúdo do modelo
    paragrafos = MODELO_OS_TEMPLATE.split('\n')
    for paragrafo in paragrafos:
        if paragrafo.strip():
            doc.add_paragraph(paragrafo)
    
    # Salva em arquivo temporário
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name

def replace_text_in_paragraph(paragraph, contexto):
    """Substitui placeholders em um único parágrafo."""
    for key, value in contexto.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(value))

def substituir_placeholders(doc, contexto):
    """Substitui os placeholders em todo o documento (parágrafos e tabelas)."""
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, contexto)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, contexto)

def gerar_os(funcionario, df_pgr, riscos_selecionados, epis_manuais, medicoes_manuais, perigo_manual, danos_manuais, categoria_manual, logo_path=None):
    """Gera uma única Ordem de Serviço para um funcionário."""
    # Usa o modelo incorporado
    modelo_path = criar_modelo_os_temporario()
    doc = Document(modelo_path)

    if logo_path:
        try:
            header_table = doc.tables[0]
            cell = header_table.cell(0, 0)
            cell.text = "" 
            p = cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(logo_path, width=Inches(2.0))
        except IndexError:
            st.warning("Aviso: Não foi encontrada uma tabela de cabeçalho no modelo para inserir a logo.")
            p = doc.paragraphs[0]
            run = p.insert_paragraph_before().add_run()
            run.add_picture(logo_path, width=Inches(2.0))

    riscos_info = df_pgr[df_pgr['risco'].isin(riscos_selecionados)]
    
    riscos_por_categoria = {"fisico": [], "quimico": [], "biologico": [], "ergonomico": [], "acidente": []}
    danos_por_categoria = {"fisico": [], "quimico": [], "biologico": [], "ergonomico": [], "acidente": []}
    epis_recomendados = set()
    
    for _, risco_row in riscos_info.iterrows():
        categoria = str(risco_row.get("categoria", "")).strip().lower()
        risco_nome = str(risco_row.get("risco", "")).strip()
        if categoria in riscos_por_categoria:
            riscos_por_categoria[categoria].append(risco_nome)
            danos = risco_row.get("possiveis_danos")
            epis = risco_row.get("epis_recomendados")
            if pd.notna(danos): danos_por_categoria[categoria].append(str(danos))
            if pd.notna(epis): epis_recomendados.update([epi.strip() for epi in str(epis).split(',')])

    if perigo_manual and categoria_manual:
        map_categorias = {"Acidentes": "acidente", "Biológicos": "biologico", "Ergonômicos": "ergonomico", "Físicos": "fisico", "Químicos": "quimico"}
        categoria_alvo = map_categorias.get(categoria_manual)
        if categoria_alvo:
            riscos_por_categoria[categoria_alvo].append(perigo_manual)
            if danos_manuais:
                danos_por_categoria[categoria_alvo].append(danos_manuais)

    if epis_manuais:
        epis_extras = [epi.strip() for epi in epis_manuais.split(',')]
        epis_recomendados.update(epis_extras)

    medicoes_lista = []
    if 'medicoes' in df_pgr.columns:
        medicoes_df = df_pgr[df_pgr['risco'].isin(riscos_selecionados)]
        medicoes_lista = [f"{row['risco']}: {row['medicoes']}" for _, row in medicoes_df.iterrows() if 'medicoes' in row and pd.notna(row['medicoes'])]

    if medicoes_manuais:
        medicoes_lista.extend([med.strip() for med in medicoes_manuais.split('\n') if med.strip()])

    # Correção: Tratamento para data de admissão ausente
    data_admissao = "não informado"
    if 'data_de_admissao' in funcionario and pd.notna(funcionario['data_de_admissao']):
        try:
            data_admissao = pd.to_datetime(funcionario['data_de_admissao']).strftime('%d/%m/%Y')
        except Exception:
            data_admissao = str(funcionario['data_de_admissao'])

    # Correção: Remover colchetes do nome dos funcionários
    nome_funcionario = str(funcionario.get("nome_do_funcionario", "N/A")).replace("[", "").replace("]", "")

    # Nova funcionalidade: Pegar descrição de atividades da planilha
    descricao_atividades = "Não informado"
    if 'descricao_de_atividades' in funcionario and pd.notna(funcionario['descricao_de_atividades']):
        descricao_atividades = str(funcionario['descricao_de_atividades'])

    # Correção: Tratamento para riscos não selecionados
    def tratar_risco_vazio(lista_riscos):
        if not lista_riscos or all(not r.strip() for r in lista_riscos):
            return "Não identificado no momento da avaliação"
        return ", ".join(lista_riscos)

    def tratar_danos_vazios(lista_danos):
        if not lista_danos or all(not d.strip() for d in lista_danos):
            return "Não identificado no momento da avaliação"
        return "; ".join(set(lista_danos))

    contexto = {
        "[NOME EMPRESA]": str(funcionario.get("empresa", "N/A")), 
        "[UNIDADE]": str(funcionario.get("unidade", "N/A")),
        "[NOME FUNCIONÁRIO]": nome_funcionario, 
        "[DATA DE ADMISSÃO]": data_admissao,
        "[SETOR]": str(funcionario.get("setor", "N/A")), 
        "[FUNÇÃO]": str(funcionario.get("funcao", "N/A")),
        "[DESCRIÇÃO DE ATIVIDADES]": descricao_atividades,
        "[RISCOS FÍSICOS]": tratar_risco_vazio(riscos_por_categoria["fisico"]),
        "[RISCOS DE ACIDENTE]": tratar_risco_vazio(riscos_por_categoria["acidente"]),
        "[RISCOS QUÍMICOS]": tratar_risco_vazio(riscos_por_categoria["quimico"]),
        "[RISCOS BIOLÓGICOS]": tratar_risco_vazio(riscos_por_categoria["biologico"]),
        "[RISCOS ERGONÔMICOS]": tratar_risco_vazio(riscos_por_categoria["ergonomico"]),
        "[POSSÍVEIS DANOS RISCOS FÍSICOS]": tratar_danos_vazios(danos_por_categoria["fisico"]),
        "[POSSÍVEIS DANOS RISCOS ACIDENTE]": tratar_danos_vazios(danos_por_categoria["acidente"]),
        "[POSSÍVEIS DANOS RISCOS QUÍMICOS]": tratar_danos_vazios(danos_por_categoria["quimico"]),
        "[POSSÍVEIS DANOS RISCOS BIOLÓGICOS]": tratar_danos_vazios(danos_por_categoria["biologico"]),
        "[POSSÍVEIS DANOS RISCOS ERGONÔMICOS]": tratar_danos_vazios(danos_por_categoria["ergonomico"]),
        "[EPIS]": ", ".join(sorted(list(epis_recomendados))) or "Nenhum",
        "[MEDIÇÕES]": "\n".join(medicoes_lista) or "Nenhuma medição aplicável.",
    }
    
    substituir_placeholders(doc, contexto)
    
    # Remove o arquivo temporário
    os.unlink(modelo_path)
    
    return doc

# --- Base de dados PGR incorporada ---
def obter_dados_pgr():
    """Retorna os dados PGR padrão incorporados no sistema."""
    return pd.DataFrame([
        {'categoria': 'fisico', 'risco': 'Ruído', 'possiveis_danos': 'Perda auditiva, zumbido, estresse, irritabilidade.'},
        {'categoria': 'fisico', 'risco': 'Vibração', 'possiveis_danos': 'Doenças osteomusculares, problemas circulatórios.'},
        {'categoria': 'fisico', 'risco': 'Calor', 'possiveis_danos': 'Desidratação, insolação, cãibras, exaustão, intermação.'},
        {'categoria': 'fisico', 'risco': 'Frio', 'possiveis_danos': 'Hipotermia, congelamento, doenças respiratórias.'},
        {'categoria': 'fisico', 'risco': 'Radiações Ionizantes', 'possiveis_danos': 'Câncer, mutações genéticas, queimaduras.'},
        {'categoria': 'fisico', 'risco': 'Radiações Não Ionizantes', 'possiveis_danos': 'Queimaduras, lesões oculares, câncer de pele.'},
        {'categoria': 'fisico', 'risco': 'Pressões Anormais', 'possiveis_danos': 'Doença descompressiva, barotrauma.'},
        {'categoria': 'fisico', 'risco': 'Umidade', 'possiveis_danos': 'Doenças respiratórias, dermatites.'},
        {'categoria': 'quimico', 'risco': 'Poeiras', 'possiveis_danos': 'Pneumoconioses, irritação respiratória, alergias.'},
        {'categoria': 'quimico', 'risco': 'Fumos', 'possiveis_danos': 'Doenças respiratórias, intoxicações.'},
        {'categoria': 'quimico', 'risco': 'Névoas', 'possiveis_danos': 'Irritação respiratória, dermatites.'},
        {'categoria': 'quimico', 'risco': 'Gases', 'possiveis_danos': 'Asfixia, intoxicações, irritação respiratória.'},
        {'categoria': 'quimico', 'risco': 'Vapores', 'possiveis_danos': 'Irritação respiratória, intoxicações, dermatites.'},
        {'categoria': 'quimico', 'risco': 'Substâncias Químicas (líquidos e sólidos)', 'possiveis_danos': 'Queimaduras, irritações, intoxicações, dermatites, câncer.'},
        {'categoria': 'quimico', 'risco': 'Agrotóxicos', 'possiveis_danos': 'Intoxicações, dermatites, câncer.'},
        {'categoria': 'biologico', 'risco': 'Bactérias', 'possiveis_danos': 'Infecções, doenças infecciosas.'},
        {'categoria': 'biologico', 'risco': 'Fungos', 'possiveis_danos': 'Micoses, alergias, infecções respiratórias.'},
        {'categoria': 'biologico', 'risco': 'Vírus', 'possiveis_danos': 'Doenças virais, infecções.'},
        {'categoria': 'biologico', 'risco': 'Parasitas', 'possiveis_danos': 'Doenças parasitárias, infecções.'},
        {'categoria': 'biologico', 'risco': 'Protozoários', 'possiveis_danos': 'Doenças parasitárias.'},
        {'categoria': 'biologico', 'risco': 'Parasitas e Protozoários', 'possiveis_danos': 'Doenças parasitárias.'},
        {'categoria': 'ergonomico', 'risco': 'Levantamento e Transporte Manual de Peso', 'possiveis_danos': 'Lesões musculoesqueléticas, dores na coluna.'},
        {'categoria': 'ergonomico', 'risco': 'Posturas Inadequadas', 'possiveis_danos': 'Dores musculares, lesões na coluna, LER/DORT.'},
        {'categoria': 'ergonomico', 'risco': 'Repetitividade', 'possiveis_danos': 'LER/DORT, tendinites, síndrome do túnel do carpo.'},
        {'categoria': 'ergonomico', 'risco': 'Jornada de Trabalho Prolongada', 'possiveis_danos': 'Fadiga, estresse, acidentes de trabalho.'},
        {'categoria': 'ergonomico', 'risco': 'Monotonia e Ritmo Excessivo', 'possiveis_danos': 'Estresse, fadiga mental, desmotivação.'},
        {'categoria': 'ergonomico', 'risco': 'Controle Rígido de Produtividade', 'possiveis_danos': 'Estresse, ansiedade, burnout.'},
        {'categoria': 'ergonomico', 'risco': 'Iluminação Inadequada', 'possiveis_danos': 'Fadiga visual, dores de cabeça.'},
        {'categoria': 'ergonomico', 'risco': 'Mobiliário Inadequado', 'possiveis_danos': 'Dores musculares, lesões na coluna.'},
        {'categoria': 'ergonomico', 'risco': 'Ritmo de Trabalho Excessivo', 'possiveis_danos': 'Estresse, fadiga, LER/DORT.'},
        {'categoria': 'acidente', 'risco': 'Arranjo Físico Inadequado', 'possiveis_danos': 'Quedas, colisões, esmagamentos.'},
        {'categoria': 'acidente', 'risco': 'Máquinas e Equipamentos sem Proteção', 'possiveis_danos': 'Amputações, cortes, esmagamentos, prensamentos.'},
        {'categoria': 'acidente', 'risco': 'Ferramentas Inadequadas ou Defeituosas', 'possiveis_danos': 'Cortes, perfurações, fraturas.'},
        {'categoria': 'acidente', 'risco': 'Eletricidade', 'possiveis_danos': 'Choque elétrico, queimaduras, fibrilação ventricular.'},
        {'categoria': 'acidente', 'risco': 'Incêndio e Explosão', 'possiveis_danos': 'Queimaduras, asfixia, lesões por impacto.'},
        {'categoria': 'acidente', 'risco': 'Animais Peçonhentos', 'possiveis_danos': 'Picadas, mordidas, reações alérgicas, envenenamento.'},
        {'categoria': 'acidente', 'risco': 'Armazenamento Inadequado', 'possiveis_danos': 'Quedas de materiais, esmagamentos, soterramentos.'},
        {'categoria': 'acidente', 'risco': 'Trabalho em Altura', 'possiveis_danos': 'Quedas, fraturas, morte.'},
        {'categoria': 'acidente', 'risco': 'Espaços Confinados', 'possiveis_danos': 'Asfixia, intoxicações, explosões.'},
        {'categoria': 'acidente', 'risco': 'Condução de Veículos', 'possiveis_danos': 'Acidentes de trânsito, lesões diversas.'},
        {'categoria': 'acidente', 'risco': 'Outros (especificar)', 'possiveis_danos': 'Variados, dependendo do risco específico.'}
    ])

# --- Interface do Streamlit ---

# Estilo CSS para layout mais profissional e correção de visibilidade
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #1e3a8a 0%, #3b82f6 100%);
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
        color: white;
    }
    .section-header {
        background: #f8fafc;
        padding: 1rem;
        border-left: 4px solid #3b82f6;
        margin: 1rem 0;
        border-radius: 5px;
    }
    .info-box {
        background: #e0f2fe;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #0284c7;
        margin: 1rem 0;
    }
    .success-box {
        background: #dcfce7;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #16a34a;
        margin: 1rem 0;
    }
    .stButton > button {
        background: linear-gradient(90deg, #1e3a8a 0%, #3b82f6 100%);
        color: white;
        border: none;
        padding: 0.5rem 2rem;
        border-radius: 8px;
        font-weight: 600;
    }
    .stSelectbox > div > div {
        background-color: #f8fafc;
    }
    /* Correção para visibilidade do texto nas tabelas */
    .stDataFrame {
        background-color: white;
    }
    .stDataFrame table {
        background-color: white !important;
        color: black !important; /* Força a cor do texto para preto */
    }
    .stDataFrame th {
        background-color: #f9fafb !important;
        color: black !important; /* Força a cor do texto para preto */
        font-weight: 600;
    }
    .stDataFrame td {
        background-color: white !important;
        color: black !important; /* Força a cor do texto para preto */
    }
    /* Correção para texto em elementos de entrada */
    .stTextInput > div > div > input {
        color: black !important; /* Força a cor do texto para preto */
        background-color: white !important;
    }
    .stTextArea > div > div > textarea {
        color: black !important; /* Força a cor do texto para preto */
        background-color: white !important;
    }
    .stSelectbox > div > div > div {
        color: black !important; /* Força a cor do texto para preto */
        background-color: white !important;
    }
</style>
""", unsafe_allow_html=True)

# Cabeçalho principal
st.markdown("""
<div class="main-header">
    <h1>📄 Gerador de Ordens de Serviço (OS)</h1>
    <p>Sistema profissional para geração automática de Ordens de Serviço de Segurança e Saúde no Trabalho</p>
</div>
""", unsafe_allow_html=True)

# --- Sidebar para upload de arquivos ---
st.sidebar.markdown("### 📁 Carregar Arquivos")
arquivo_funcionarios = st.sidebar.file_uploader(
    "Planilha de Funcionários", 
    type="xlsx", 
    help="Ficheiro .xlsx obrigatório com os dados dos funcionários."
)

arquivo_logo = st.sidebar.file_uploader(
    "Logo da Empresa (Opcional)", 
    type=["png", "jpg", "jpeg"],
    help="Imagem da logo que será inserida no cabeçalho da OS"
)

st.sidebar.markdown("---")
st.sidebar.markdown("### ℹ️ Informações")
st.sidebar.info("O modelo de OS já está incorporado no sistema. Não é necessário fazer upload do modelo.")

# --- Carregamento e Processamento dos Dados ---
df_pgr = obter_dados_pgr()

if arquivo_funcionarios:
    df_funcionarios = carregar_planilha(arquivo_funcionarios)
    if df_funcionarios is not None:
        df_funcionarios = mapear_e_renomear_colunas_funcionarios(df_funcionarios)
        
        st.markdown('<div class="section-header"><h3>👥 Seleção de Funcionários</h3></div>', unsafe_allow_html=True)
        
        # Filtros para seleção automática
        col1, col2 = st.columns(2)
        
        with col1:
            setores_disponiveis = df_funcionarios['setor'].dropna().unique().tolist() if 'setor' in df_funcionarios.columns else []
            setor_selecionado = st.selectbox("Selecionar Setor", ["Todos"] + setores_disponiveis)
        
        with col2:
            if setor_selecionado != "Todos":
                df_filtrado = df_funcionarios[df_funcionarios['setor'] == setor_selecionado]
            else:
                df_filtrado = df_funcionarios
                
            funcoes_disponiveis = df_filtrado['funcao'].dropna().unique().tolist() if 'funcao' in df_filtrado.columns else []
            funcao_selecionada = st.selectbox("Selecionar Função/Cargo", ["Todos"] + funcoes_disponiveis)
        
        # Seleção automática de funcionários baseada nos filtros
        if setor_selecionado != "Todos":
            df_funcionarios_filtrados = df_funcionarios[df_funcionarios['setor'] == setor_selecionado]
        else:
            df_funcionarios_filtrados = df_funcionarios
            
        if funcao_selecionada != "Todos":
            df_funcionarios_filtrados = df_funcionarios_filtrados[df_funcionarios_filtrados['funcao'] == funcao_selecionada]
        
        st.markdown(f'<div class="info-box">✅ <strong>{len(df_funcionarios_filtrados)} funcionários</strong> selecionados automaticamente</div>', unsafe_allow_html=True)
        
        # Mostrar TODOS os funcionários selecionados
        if len(df_funcionarios_filtrados) > 0:
            colunas_para_mostrar = ['nome_do_funcionario', 'setor', 'funcao']
            if 'descricao_de_atividades' in df_funcionarios_filtrados.columns:
                colunas_para_mostrar.append('descricao_de_atividades')
            
            st.dataframe(df_funcionarios_filtrados[colunas_para_mostrar], use_container_width=True)
        
        # --- Configuração de Riscos ---
        st.markdown('<div class="section-header"><h3>⚠️ Configuração de Riscos</h3></div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Riscos da Base de Dados")
            riscos_disponiveis = df_pgr['risco'].unique().tolist()
            riscos_selecionados = st.multiselect("Selecionar Riscos", riscos_disponiveis)
        
        with col2:
            st.subheader("Adicionar Risco Manual")
            categoria_manual = st.selectbox("Categoria do Risco", ["", "Físicos", "Químicos", "Biológicos", "Ergonômicos", "Acidentes"])
            perigo_manual = st.text_input("Descrição do Risco")
            danos_manuais = st.text_area("Possíveis Danos", placeholder="Descreva os possíveis danos...")
        
        # --- Configuração de EPIs ---
        st.markdown('<div class="section-header"><h3>🦺 Equipamentos de Proteção Individual (EPIs)</h3></div>', unsafe_allow_html=True)
        epis_manuais = st.text_area("EPIs Recomendados", placeholder="Digite os EPIs separados por vírgula...")
        
        # --- Configuração de Medições ---
        st.markdown('<div class="section-header"><h3>📊 Medições e Avaliações</h3></div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            medicao_valor = st.text_input("Valor da Medição")
        with col2:
            unidade_medicao = st.selectbox("Unidade de Medida", UNIDADES_MEDIDA)
        
        medicao_descricao = st.text_area("Descrição da Medição", placeholder="Descreva o contexto da medição...")
        
        medicoes_manuais = ""
        if medicao_valor and unidade_medicao:
            medicoes_manuais = f"{medicao_descricao}: {medicao_valor} {unidade_medicao}"
        
        # --- Geração das OSs ---
        st.markdown('<div class="section-header"><h3>🚀 Gerar Ordens de Serviço</h3></div>', unsafe_allow_html=True)
        
        if st.button("🔄 Gerar OSs para Funcionários Selecionados", type="primary"):
            if len(df_funcionarios_filtrados) == 0:
                st.error("Nenhum funcionário selecionado!")
            else:
                with st.spinner("Gerando Ordens de Serviço..."):
                    documentos_gerados = []
                    logo_path = None
                    
                    if arquivo_logo:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{arquivo_logo.name.split('.')[-1]}") as temp_logo:
                            temp_logo.write(arquivo_logo.read())
                            logo_path = temp_logo.name
                    
                    progress_bar = st.progress(0)
                    total_funcionarios = len(df_funcionarios_filtrados)
                    
                    for i, (_, funcionario) in enumerate(df_funcionarios_filtrados.iterrows()):
                        try:
                            doc = gerar_os(
                                funcionario, df_pgr, 
                                riscos_selecionados, epis_manuais, medicoes_manuais,
                                perigo_manual, danos_manuais, categoria_manual, logo_path
                            )
                            
                            # Salva o documento em memória
                            doc_io = BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)
                            
                            nome_limpo = str(funcionario.get("nome_do_funcionario", f"Funcionario_{i+1}")).replace("[", "").replace("]", "")
                            documentos_gerados.append((f"OS_{nome_limpo.replace(' ', '_')}.docx", doc_io.getvalue()))
                            
                            progress_bar.progress((i + 1) / total_funcionarios)
                            
                        except Exception as e:
                            st.error(f"Erro ao gerar OS para {funcionario.get('nome_do_funcionario', 'funcionário')}: {e}")
                    
                    if logo_path:
                        os.unlink(logo_path)
                    
                    if documentos_gerados:
                        # Criar ZIP com todos os documentos
                        zip_buffer = BytesIO()
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            for nome_arquivo, conteudo in documentos_gerados:
                                zip_file.writestr(nome_arquivo, conteudo)
                        
                        zip_buffer.seek(0)
                        
                        st.markdown(f'<div class="success-box">✅ <strong>{len(documentos_gerados)} Ordens de Serviço</strong> geradas com sucesso!</div>', unsafe_allow_html=True)
                        
                        # Download automático
                        st.download_button(
                            label="📥 Baixar Todas as OSs (ZIP)",
                            data=zip_buffer.getvalue(),
                            file_name=f"Ordens_de_Servico_{time.strftime('%Y%m%d_%H%M%S')}.zip",
                            mime="application/zip",
                            key="download_zip"
                        )
                        
                        # Trigger download automático via JavaScript
                        st.markdown("""
                        <script>
                        setTimeout(function(){
                            document.querySelector('[data-testid="stDownloadButton"] button').click();
                        }, 1000);
                        </script>
                        """, unsafe_allow_html=True)
else:
    st.markdown('<div class="info-box">📋 Por favor, carregue a planilha de funcionários para começar.</div>', unsafe_allow_html=True)
    
    # Exemplo de estrutura da planilha
    st.markdown('<div class="section-header"><h3>📋 Estrutura da Planilha de Funcionários</h3></div>', unsafe_allow_html=True)
    exemplo_df = pd.DataFrame({
        'Nome': ['João Silva', 'Maria Santos'],
        'Setor': ['Produção', 'Administrativo'],
        'Função': ['Operador', 'Assistente'],
        'Data de Admissão': ['01/01/2020', '15/03/2021'],
        'Empresa': ['Empresa XYZ', 'Empresa XYZ'],
        'Unidade': ['Matriz', 'Filial'],
        'Descrição de Atividades': ['Operar máquinas de produção', 'Atividades administrativas gerais']
    })
    st.dataframe(exemplo_df, use_container_width=True)
    st.info("A planilha deve conter pelo menos as colunas: Nome, Setor, Função. A coluna 'Descrição de Atividades' é opcional mas recomendada.")

