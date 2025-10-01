import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile
from io import BytesIO
import time
import re
from datetime import datetime
import sys
import os

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from database.models import DatabaseManager
from database.auth import AuthManager
from database.user_data import UserDataManager

# Configuração
st.set_page_config(
    page_title="Gerador de OS",
    page_icon="📋",
    layout="wide",
)

# CSS Melhorado
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .user-card {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #667eea;
        margin-bottom: 1rem;
    }
    .metric-box {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px solid #e9ecef;
        text-align: center;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .metric-box h3 {
        color: #667eea;
        margin: 0;
        font-size: 2rem;
    }
    .metric-box p {
        color: #6c757d;
        margin: 0;
        font-size: 0.9rem;
    }
    .risk-category {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        margin-bottom: 1rem;
        border-left: 4px solid;
    }
    .risk-fisico { border-left-color: #dc3545; }
    .risk-quimico { border-left-color: #fd7e14; }
    .risk-biologico { border-left-color: #28a745; }
    .risk-ergonomico { border-left-color: #007bff; }
    .risk-acidente { border-left-color: #ffc107; }
    .risk-manual { border-left-color: #6f42c1; }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px 8px 0 0;
        padding: 10px 20px;
    }
</style>
""", unsafe_allow_html=True)

# Constantes
UNIDADES_DE_MEDIDA = ["dB(A)", "m/s²", "m/s¹·⁷⁵", "ppm", "mg/m³", "%", "°C", "lx", "cal/cm²", "µT", "kV/m", "W/m²", "f/cm³", "Não aplicável"]
CATEGORIAS_RISCO = {
    'fisico': '🔥 Físicos',
    'quimico': '⚗️ Químicos',
    'biologico': '🦠 Biológicos',
    'ergonomico': '🏃 Ergonômicos',
    'acidente': '⚠️ Acidentes'
}

# Dados PGR (mantidos conforme original)
RISCOS_PGR_DADOS = {
    'quimico': {
        'riscos': ['Exposição a Produto Químico'],
        'danos': ['Irritação/lesão ocular, na pele e mucosas; Dermatites; Queimadura Química; Intoxicação; Náuseas; Vômitos.']
    },
    'fisico': {
        'riscos': ['Ambiente Artificialmente Frio', 'Exposição ao Ruído', 'Vibrações Localizadas (mão/braço)', 
                   'Vibração de Corpo Inteiro (AREN)', 'Vibração de Corpo Inteiro (VDVR)', 'Exposição à Radiações Ionizantes',
                   'Exposição à Radiações Não-ionizantes', 'Exposição à Temperatura Ambiente Elevada', 
                   'Exposição à Temperatura Ambiente Baixa', 'Pressão Atmosférica Anormal (condições hiperbáricas)', 'Umidade'],
        'danos': ['Estresse, desconforto, dormência, rigidez nas partes com maior intensidade de exposição ao frio.',
                  'Perda Auditiva Induzida pelo Ruído Ocupacional (PAIRO).', 'Alterações articulares e vasomotoras.',
                  'Alterações no sistema digestivo, sistema musculoesquelético, sistema nervoso.', 
                  'Alterações no sistema digestivo, sistema musculoesquelético, sistema nervoso.',
                  'Dano às células do corpo humano, causando doenças graves, inclusive fatais, como câncer.',
                  'Depressão imunológica, fotoenvelhecimento, lesões oculares.', 
                  'Desidratação, erupções cutâneas, cãibras, fadiga física.',
                  'Estresse, desconforto, dormência, rigidez nas partes expostas ao frio.',
                  'Barotrauma pulmonar, lesão de tecido pulmonar ou pneumotórax.', 
                  'Doenças do aparelho respiratório, quedas, doenças de pele.']
    },
    'biologico': {
        'riscos': ['Água e/ou alimentos contaminados', 'Contato com Fluido Orgânico (sangue, hemoderivados)', 
                   'Contato com Pessoas Doentes e/ou Material Infectocontagiante', 'Contaminação pelo Corona Vírus',
                   'Exposição à Agentes Microbiológicos (fungos, bactérias, vírus)'],
        'danos': ['Intoxicação, diarreias, infecções intestinais.', 'Doenças infectocontagiosas.', 
                  'Doenças infectocontagiosas.', 'COVID-19, podendo causar gripes, febre, tosse seca.',
                  'Doenças infectocontagiosas, dermatites, irritação.']
    },
    'ergonomico': {
        'riscos': ['Posturas incômodas/pouco confortáveis por longos períodos', 'Postura sentada por longos períodos',
                   'Postura em pé por longos períodos', 'Frequente deslocamento à pé durante à jornada',
                   'Esforço físico intenso', 'Levantamento e transporte manual de cargas',
                   'Frequente execução de movimentos repetitivos', 'Uso frequente de força, pressão, preensão'],
        'danos': ['Distúrbios musculoesqueléticos em músculos e articulações dos membros superiores, inferiores e coluna.',
                  'Sobrecarga dos membros superiores e coluna vertebral; Dor localizada.',
                  'Sobrecarga corporal, dores nos membros inferiores e coluna vertebral.',
                  'Sobrecarga corporal, dores nos membros inferiores e coluna.', 
                  'Distúrbios musculoesqueléticos; Fadiga, Dor localizada.',
                  'Distúrbios musculoesqueléticos; Fadiga, Dor localizada.',
                  'Distúrbios osteomusculares em músculos e articulações dos membros.',
                  'Sobrecarga muscular, fadiga, dor localizada.']
    },
    'acidente': {
        'riscos': ['Absorção (por contato) de substância cáustica, tóxica', 'Afogamento, imersão, engolfamento',
                   'Aprisionamento em, sob ou entre', 'Ataque de ser vivo por mordedura, picada',
                   'Atrito ou abrasão por encostar em objeto', 'Atropelamento', 'Batida contra objeto parado',
                   'Carga Suspensa', 'Colisão entre veículos e/ou equipamentos', 'Exposição à Energia Elétrica',
                   'Incêndio/Explosão', 'Objetos cortantes/perfurocortantes', 'Queda de pessoa com diferença de nível'],
        'danos': ['Intoxicação, envenenamento, queimadura, irritação.', 'Asfixia, desconforto respiratório.',
                  'Compressão/esmagamento de partes do corpo, cortes, escoriações, fraturas.',
                  'Perfurações, cortes, arranhões, escoriações.', 'Cortes, ferimentos, esfoladura.',
                  'Compressão/esmagamento de partes do corpo, cortes, fraturas.', 'Cortes, escoriações, fraturas.',
                  'Esmagamento, prensamento de partes do corpo.', 'Compressão/esmagamento, cortes, fraturas.',
                  'Choque elétrico e eletroplessão (eletrocussão).', 'Queimadura de 1º, 2º ou 3º grau, asfixia.',
                  'Corte, laceração, ferida contusa, punctura (ferida aberta).', 'Escoriações, ferimentos, fraturas, morte.']
    }
}

def get_danos_por_riscos_pgr(categoria, riscos_selecionados):
    """Retorna danos associados aos riscos selecionados"""
    if categoria not in RISCOS_PGR_DADOS or not riscos_selecionados:
        return ""
    
    danos_lista = []
    riscos_categoria = RISCOS_PGR_DADOS[categoria]["riscos"]
    danos_categoria = RISCOS_PGR_DADOS[categoria]["danos"]
    
    for risco in riscos_selecionados:
        if risco in riscos_categoria:
            indice = riscos_categoria.index(risco)
            if indice < len(danos_categoria):
                danos_lista.append(danos_categoria[indice])
    
    return "; ".join(danos_lista) if danos_lista else ""

# Inicialização
@st.cache_resource
def init_managers():
    db_manager = DatabaseManager()
    auth_manager = AuthManager(db_manager)
    user_data_manager = UserDataManager(db_manager)
    return db_manager, auth_manager, user_data_manager

db_manager, auth_manager, user_data_manager = init_managers()

def show_login_page():
    st.markdown('<div class="main-header"><h1>🔐 Gerador de Ordens de Serviço</h1><p>Sistema Profissional de Gestão de OS</p></div>', unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["🔑 Login", "📝 Cadastro"])
    
    with tab1:
        with st.form("login_form"):
            email = st.text_input("📧 Email", placeholder="seu@email.com")
            password = st.text_input("🔒 Senha", type="password")
            
            if st.form_submit_button("🚀 Entrar", use_container_width=True, type="primary"):
                if email and password:
                    success, message, session_data = auth_manager.login_user(email, password)
                    if success:
                        st.session_state.authenticated = True
                        st.session_state.user_data = session_data
                        st.session_state.user_data_loaded = False
                        st.success(message)
                        time.sleep(0.5)
                        st.rerun()
                    else:
                        st.error(message)
                else:
                    st.warning("Preencha todos os campos")
    
    with tab2:
        with st.form("register_form"):
            reg_email = st.text_input("📧 Email", placeholder="seu@email.com", key="reg_email")
            reg_password = st.text_input("🔒 Senha (mín. 6 caracteres)", type="password", key="reg_password")
            reg_password_confirm = st.text_input("🔒 Confirmar Senha", type="password")
            
            if st.form_submit_button("✅ Criar Conta", use_container_width=True, type="primary"):
                if reg_email and reg_password and reg_password_confirm:
                    if reg_password != reg_password_confirm:
                        st.error("❌ As senhas não coincidem")
                    else:
                        success, message = auth_manager.register_user(reg_email, reg_password)
                        if success:
                            st.success(message)
                            st.info("✅ Agora você pode fazer login!")
                        else:
                            st.error(message)
                else:
                    st.warning("Preencha todos os campos")

def check_authentication():
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'user_data' not in st.session_state:
        st.session_state.user_data = None
    if st.session_state.authenticated and st.session_state.user_data:
        session_token = st.session_state.user_data.get('session_token')
        if session_token:
            is_valid, _ = auth_manager.validate_session(session_token)
            if not arquivo_funcionarios or not arquivo_modelo_os:
        st.info("📋 Carregue os documentos para continuar")
        return
    
    df_funcionarios_raw = carregar_planilha(arquivo_funcionarios)
    if df_funcionarios_raw is None:
        st.stop()

    df_funcionarios = mapear_e_renomear_colunas_funcionarios(df_funcionarios_raw)

    # Seleção de funcionários
    with st.container(border=True):
        st.markdown('#### 👥 2. Selecionar Funcionários')
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown('<div class="metric-box"><h3>📊</h3><p>Total de Funcionários</p><h3>' + str(len(df_funcionarios)) + '</h3></div>', unsafe_allow_html=True)
        
        with col2:
            setores = sorted(df_funcionarios['setor'].dropna().unique().tolist()) if 'setor' in df_funcionarios.columns else []
            st.markdown('<div class="metric-box"><h3>🏢</h3><p>Setores</p><h3>' + str(len(setores)) + '</h3></div>', unsafe_allow_html=True)
        
        with col3:
            funcoes = sorted(df_funcionarios['funcao'].dropna().unique().tolist()) if 'funcao' in df_funcionarios.columns else []
            st.markdown('<div class="metric-box"><h3>💼</h3><p>Funções</p><h3>' + str(len(funcoes)) + '</h3></div>', unsafe_allow_html=True)
        
        st.markdown("---")
        
        setor_sel = st.multiselect("🏢 Filtrar por Setor(es)", setores, help="Selecione um ou mais setores")
        df_filtrado_setor = df_funcionarios[df_funcionarios['setor'].isin(setor_sel)] if setor_sel else df_funcionarios
        
        funcoes_disponiveis = sorted(df_filtrado_setor['funcao'].dropna().unique().tolist()) if 'funcao' in df_filtrado_setor.columns else []
        funcoes_formatadas = []
        if setor_sel:
            for funcao in funcoes_disponiveis:
                concluido = all((s, funcao) in st.session_state.cargos_concluidos for s in setor_sel)
                if concluido:
                    funcoes_formatadas.append(f"{funcao} ✅")
                else:
                    funcoes_formatadas.append(funcao)
        else:
            funcoes_formatadas = funcoes_disponiveis
        
        funcao_sel_formatada = st.multiselect("💼 Filtrar por Função/Cargo(s)", funcoes_formatadas)
        funcao_sel = [f.replace(" ✅", "") for f in funcao_sel_formatada]
        
        df_final_filtrado = df_filtrado_setor[df_filtrado_setor['funcao'].isin(funcao_sel)] if funcao_sel else df_filtrado_setor
        
        if len(df_final_filtrado) > 0:
            st.success(f"✅ **{len(df_final_filtrado)} funcionário(s) selecionado(s)**")
            with st.expander("👀 Visualizar funcionários selecionados"):
                st.dataframe(df_final_filtrado[['nome_do_funcionario', 'setor', 'funcao']], use_container_width=True)
        else:
            st.warning("⚠️ Nenhum funcionário selecionado")

    # Configuração de riscos
    with st.container(border=True):
        st.markdown('#### ⚠️ 3. Configurar Riscos Ocupacionais')
        st.info("💡 Configure os riscos que serão aplicados a TODOS os funcionários selecionados")

        tab_fisico, tab_quimico, tab_biologico, tab_ergonomico, tab_acidente, tab_manual = st.tabs([
            "🔥 Físicos", "⚗️ Químicos", "🦠 Biológicos", "🏃 Ergonômicos", "⚠️ Acidentes", "➕ Manual"
        ])

        riscos_selecionados_pgr = {}

        with tab_fisico:
            st.markdown('<div class="risk-category risk-fisico"><h4>🔥 Riscos Físicos</h4></div>', unsafe_allow_html=True)
            if 'fisico' in RISCOS_PGR_DADOS:
                st.caption(f"📋 {len(RISCOS_PGR_DADOS['fisico']['riscos'])} opções disponíveis")
                riscos_selecionados_pgr['fisico'] = st.multiselect(
                    "Selecione os riscos:",
                    options=RISCOS_PGR_DADOS['fisico']['riscos'],
                    key="riscos_pgr_fisico"
                )
                if riscos_selecionados_pgr['fisico']:
                    danos = get_danos_por_riscos_pgr('fisico', riscos_selecionados_pgr['fisico'])
                    if danos:
                        st.info(f"**Possíveis Danos:** {danos}")

        with tab_quimico:
            st.markdown('<div class="risk-category risk-quimico"><h4>⚗️ Riscos Químicos</h4></div>', unsafe_allow_html=True)
            if 'quimico' in RISCOS_PGR_DADOS:
                st.caption(f"📋 {len(RISCOS_PGR_DADOS['quimico']['riscos'])} opções disponíveis")
                riscos_selecionados_pgr['quimico'] = st.multiselect(
                    "Selecione os riscos:",
                    options=RISCOS_PGR_DADOS['quimico']['riscos'],
                    key="riscos_pgr_quimico"
                )
                if riscos_selecionados_pgr['quimico']:
                    danos = get_danos_por_riscos_pgr('quimico', riscos_selecionados_pgr['quimico'])
                    if danos:
                        st.info(f"**Possíveis Danos:** {danos}")

        with tab_biologico:
            st.markdown('<div class="risk-category risk-biologico"><h4>🦠 Riscos Biológicos</h4></div>', unsafe_allow_html=True)
            if 'biologico' in RISCOS_PGR_DADOS:
                st.caption(f"📋 {len(RISCOS_PGR_DADOS['biologico']['riscos'])} opções disponíveis")
                riscos_selecionados_pgr['biologico'] = st.multiselect(
                    "Selecione os riscos:",
                    options=RISCOS_PGR_DADOS['biologico']['riscos'],
                    key="riscos_pgr_biologico"
                )
                if riscos_selecionados_pgr['biologico']:
                    danos = get_danos_por_riscos_pgr('biologico', riscos_selecionados_pgr['biologico'])
                    if danos:
                        st.info(f"**Possíveis Danos:** {danos}")

        with tab_ergonomico:
            st.markdown('<div class="risk-category risk-ergonomico"><h4>🏃 Riscos Ergonômicos</h4></div>', unsafe_allow_html=True)
            if 'ergonomico' in RISCOS_PGR_DADOS:
                st.caption(f"📋 {len(RISCOS_PGR_DADOS['ergonomico']['riscos'])} opções disponíveis")
                riscos_selecionados_pgr['ergonomico'] = st.multiselect(
                    "Selecione os riscos:",
                    options=RISCOS_PGR_DADOS['ergonomico']['riscos'],
                    key="riscos_pgr_ergonomico"
                )
                if riscos_selecionados_pgr['ergonomico']:
                    danos = get_danos_por_riscos_pgr('ergonomico', riscos_selecionados_pgr['ergonomico'])
                    if danos:
                        st.info(f"**Possíveis Danos:** {danos}")

        with tab_acidente:
            st.markdown('<div class="risk-category risk-acidente"><h4>⚠️ Riscos de Acidente</h4></div>', unsafe_allow_html=True)
            if 'acidente' in RISCOS_PGR_DADOS:
                st.caption(f"📋 {len(RISCOS_PGR_DADOS['acidente']['riscos'])} opções disponíveis")
                riscos_selecionados_pgr['acidente'] = st.multiselect(
                    "Selecione os riscos:",
                    options=RISCOS_PGR_DADOS['acidente']['riscos'],
                    key="riscos_pgr_acidente"
                )
                if riscos_selecionados_pgr['acidente']:
                    danos = get_danos_por_riscos_pgr('acidente', riscos_selecionados_pgr['acidente'])
                    if danos:
                        st.info(f"**Possíveis Danos:** {danos}")

        with tab_manual:
            st.markdown('<div class="risk-category risk-manual"><h4>➕ Riscos Personalizados</h4></div>', unsafe_allow_html=True)
            with st.form("form_risco_manual", clear_on_submit=True):
                risco_manual_nome = st.text_input("📝 Descrição do Risco")
                categoria_manual = st.selectbox("🏷️ Categoria", list(CATEGORIAS_RISCO.values()))
                danos_manuais = st.text_area("⚕️ Possíveis Danos (Opcional)")
                
                if st.form_submit_button("➕ Adicionar Risco", type="primary", use_container_width=True):
                    if risco_manual_nome and categoria_manual:
                        user_data_manager.add_manual_risk(user_id, categoria_manual, risco_manual_nome, danos_manuais)
                        st.session_state.user_data_loaded = False
                        st.success("✅ Risco adicionado!")
                        time.sleep(0.5)
                        st.rerun()
            
            if st.session_state.riscos_manuais_adicionados:
                st.markdown("**📋 Riscos manuais salvos:**")
                for r in st.session_state.riscos_manuais_adicionados:
                    col1, col2 = st.columns([5, 1])
                    col1.markdown(f"• **{r['risk_name']}** ({r['category']})")
                    if col2.button("🗑️", key=f"rem_risco_{r['id']}", help="Remover"):
                        user_data_manager.remove_manual_risk(user_id, r['id'])
                        st.session_state.user_data_loaded = False
                        st.rerun()

        # Resumo de riscos
        total_riscos_pgr = sum(len(riscos) for riscos in riscos_selecionados_pgr.values())
        total_riscos = total_riscos_pgr + len(st.session_state.riscos_manuais_adicionados)
        
        if total_riscos > 0:
            with st.expander(f"📊 Resumo de Riscos Selecionados ({total_riscos} total)", expanded=True):
                for categoria_key, categoria_display in CATEGORIAS_RISCO.items():
                    riscos_lista = []
                    if categoria_key in riscos_selecionados_pgr and riscos_selecionados_pgr[categoria_key]:
                        riscos_lista.extend(riscos_selecionados_pgr[categoria_key])
                    
                    for risco_manual in st.session_state.riscos_manuais_adicionados:
                        if risco_manual['category'] == categoria_display:
                            riscos_lista.append(f"{risco_manual['risk_name']} (Manual)")
                    
                    if riscos_lista:
                        st.markdown(f"**{categoria_display}** ({len(riscos_lista)})")
                        for risco in sorted(list(set(riscos_lista))):
                            st.markdown(f"• {risco}")

    # Medições e EPIs
    with st.container(border=True):
        st.markdown('#### 🛡️ 4. Medições e EPIs')
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**📊 Medições Ambientais**")
            with st.form("form_medicao", clear_on_submit=True):
                agente = st.text_input("🔬 Agente/Fonte")
                valor = st.text_input("📈 Valor Medido")
                unidade = st.selectbox("📏 Unidade", UNIDADES_DE_MEDIDA)
                epi_med = st.text_input("🦺 EPI Associado (Opcional)")
                
                if st.form_submit_button("➕ Adicionar", use_container_width=True):
                    if agente and valor:
                        user_data_manager.add_measurement(user_id, agente, valor, unidade, epi_med)
                        st.session_state.user_data_loaded = False
                        st.success("✅ Medição adicionada!")
                        time.sleep(0.5)
                        st.rerun()
            
            if st.session_state.medicoes_adicionadas:
                st.markdown("**📋 Medições salvas:**")
                for med in st.session_state.medicoes_adicionadas:
                    col_a, col_b = st.columns([5, 1])
                    col_a.markdown(f"• {med['agent']}: {med['value']} {med['unit']}")
                    if col_b.button("🗑️", key=f"rem_med_{med['id']}", help="Remover"):
                        user_data_manager.remove_measurement(user_id, med['id'])
                        st.session_state.user_data_loaded = False
                        st.rerun()
        
        with col2:
            st.markdown("**🦺 EPIs Gerais**")
            with st.form("form_epi", clear_on_submit=True):
                epi_nome = st.text_input("🛡️ Nome do EPI")
                
                if st.form_submit_button("➕ Adicionar", use_container_width=True):
                    if epi_nome:
                        user_data_manager.add_epi(user_id, epi_nome)
                        st.session_state.user_data_loaded = False
                        st.success("✅ EPI adicionado!")
                        time.sleep(0.5)
                        st.rerun()
            
            if st.session_state.epis_adicionados:
                st.markdown("**📋 EPIs salvos:**")
                for epi in st.session_state.epis_adicionados:
                    col_a, col_b = st.columns([5, 1])
                    col_a.markdown(f"• {epi['epi_name']}")
                    if col_b.button("🗑️", key=f"rem_epi_{epi['id']}", help="Remover"):
                        user_data_manager.remove_epi(user_id, epi['id'])
                        st.session_state.user_data_loaded = False
                        st.rerun()

    # Botão de geração
    st.markdown("---")
    
    if len(df_final_filtrado) == 0:
        st.warning("⚠️ Selecione funcionários para gerar as OS")
    elif total_riscos == 0:
        st.warning("⚠️ Configure pelo menos um risco ocupacional")
    else:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button(
                f"🚀 GERAR {len(df_final_filtrado)} ORDEM(S) DE SERVIÇO",
                type="primary",
                use_container_width=True
            ):
                with st.spinner(f"⏳ Gerando {len(df_final_filtrado)} documentos..."):
                    documentos_gerados = []
                    combinacoes_processadas = set()
                    
                    progress_bar = st.progress(0)
                    status = st.empty()
                    
                    for idx, (_, func) in enumerate(df_final_filtrado.iterrows()):
                        status.text(f"Processando: {func.get('nome_do_funcionario', 'N/A')} ({idx+1}/{len(df_final_filtrado)})")
                        
                        combinacoes_processadas.add((func['setor'], func['funcao']))
                        
                        # CORREÇÃO: Passar riscos_selecionados_pgr (dict) ao invés de lista
                        doc = gerar_os(
                            func,
                            riscos_selecionados_pgr,  # Dict por categoria
                            st.session_state.epis_adicionados,
                            st.session_state.medicoes_adicionadas,
                            st.session_state.riscos_manuais_adicionados,
                            arquivo_modelo_os
                        )
                        
                        doc_io = BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        nome_limpo = re.sub(r'[^\w\s-]', '', func.get("nome_do_funcionario", "Func")).strip().replace(" ", "_")
                        caminho = f"{func.get('setor', 'Setor')}/{func.get('funcao', 'Funcao')}/OS_{nome_limpo}.docx"
                        documentos_gerados.append((caminho, doc_io.getvalue()))
                        
                        progress_bar.progress((idx + 1) / len(df_final_filtrado))
                    
                    st.session_state.cargos_concluidos.update(combinacoes_processadas)
                    
                    status.empty()
                    progress_bar.empty()
                    
                    if documentos_gerados:
                        zip_buffer = BytesIO()
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            for nome_arquivo, conteudo in documentos_gerados:
                                zip_file.writestr(nome_arquivo, conteudo)
                        
                        nome_zip = f"OS_Geradas_{time.strftime('%Y%m%d_%H%M%S')}.zip"
                        
                        st.success(f"🎉 **{len(documentos_gerados)} Ordens de Serviço geradas com sucesso!**")
                        st.balloons()
                        
                        st.download_button(
                            label="📥 BAIXAR TODAS AS OS (.zip)",
                            data=zip_buffer.getvalue(),
                            file_name=nome_zip,
                            mime="application/zip",
                            type="primary",
                            use_container_width=True
                        )

if __name__ == "__main__":
    main() is_valid:
                st.session_state.authenticated = False
                st.session_state.user_data = None
                st.rerun()

def logout_user():
    if st.session_state.user_data and st.session_state.user_data.get('session_token'):
        auth_manager.logout_user(st.session_state.user_data['session_token'])
    st.session_state.authenticated = False
    st.session_state.user_data = None
    st.session_state.user_data_loaded = False
    st.rerun()

def show_user_info():
    if st.session_state.get('authenticated'):
        user_email = st.session_state.user_data.get('email', 'N/A')
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f'<div class="user-card">👤 <strong>{user_email}</strong></div>', unsafe_allow_html=True)
        with col2:
            if st.button("🚪 Sair", type="secondary", use_container_width=True):
                logout_user()

def init_user_session_state():
    if st.session_state.get('authenticated') and not st.session_state.get('user_data_loaded'):
        user_id = st.session_state.user_data.get('user_id')
        if user_id:
            st.session_state.medicoes_adicionadas = user_data_manager.get_user_measurements(user_id)
            st.session_state.epis_adicionados = user_data_manager.get_user_epis(user_id)
            st.session_state.riscos_manuais_adicionados = user_data_manager.get_user_manual_risks(user_id)
            st.session_state.user_data_loaded = True
    
    if 'medicoes_adicionadas' not in st.session_state:
        st.session_state.medicoes_adicionadas = []
    if 'epis_adicionados' not in st.session_state:
        st.session_state.epis_adicionados = []
    if 'riscos_manuais_adicionados' not in st.session_state:
        st.session_state.riscos_manuais_adicionados = []
    if 'cargos_concluidos' not in st.session_state:
        st.session_state.cargos_concluidos = set()

def normalizar_texto(texto):
    if not isinstance(texto, str): return ""
    return re.sub(r'[\s\W_]+', '', texto.lower().strip())

def mapear_e_renomear_colunas_funcionarios(df):
    df_copia = df.copy()
    mapeamento = {
        'nome_do_funcionario': ['nomedofuncionario', 'nome', 'funcionario', 'funcionário'],
        'funcao': ['funcao', 'função', 'cargo'],
        'data_de_admissao': ['datadeadmissao', 'dataadmissao', 'admissao', 'admissão'],
        'setor': ['setordetrabalho', 'setor', 'area', 'área'],
        'descricao_de_atividades': ['descricaodeatividades', 'atividades', 'descriçãodeatividades'],
        'empresa': ['empresa'],
        'unidade': ['unidade']
    }
    colunas_renomeadas = {}
    colunas_df_normalizadas = {normalizar_texto(col): col for col in df_copia.columns}
    for nome_padrao, nomes_possiveis in mapeamento.items():
        for nome_possivel in nomes_possiveis:
            if nome_possivel in colunas_df_normalizadas:
                coluna_original = colunas_df_normalizadas[nome_possivel]
                colunas_renomeadas[coluna_original] = nome_padrao
                break
    df_copia.rename(columns=colunas_renomeadas, inplace=True)
    return df_copia

@st.cache_data
def carregar_planilha(arquivo):
    if arquivo is None: return None
    try:
        return pd.read_excel(arquivo)
    except Exception as e:
        st.error(f"Erro ao ler arquivo Excel: {e}")
        return None

def substituir_placeholders(doc, contexto):
    """Substitui placeholders preservando formatação"""
    def aplicar_formatacao_padrao(run):
        run.font.name = 'Segoe UI'
        run.font.size = Pt(9)
        return run

    def processar_paragrafo(p):
        texto_original = p.text
        
        if "[MEDIÇÕES]" in texto_original:
            for run in p.runs:
                run.text = ''
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            medicoes_valor = contexto.get("[MEDIÇÕES]", "Não aplicável")
            if medicoes_valor == "Não aplicável" or not medicoes_valor.strip():
                run = aplicar_formatacao_padrao(p.add_run("Não aplicável"))
                run.font.bold = False
            else:
                linhas = medicoes_valor.split('\n')
                for i, linha in enumerate(linhas):
                    if not linha.strip(): continue
                    if i > 0: p.add_run().add_break()
                    if ":" in linha:
                        partes = linha.split(":", 1)
                        agente = partes[0].strip() + ":"
                        valor = partes[1].strip()
                        run_agente = aplicar_formatacao_padrao(p.add_run(agente + " "))
                        run_agente.font.bold = True
                        run_valor = aplicar_formatacao_padrao(p.add_run(valor))
                        run_valor.font.bold = False
                    else:
                        run = aplicar_formatacao_padrao(p.add_run(linha))
                        run.font.bold = False
            return

        placeholders_no_paragrafo = [key for key in contexto if key in texto_original]
        if not placeholders_no_paragrafo:
            return

        estilo_rotulo = {
            'bold': p.runs[0].bold if p.runs else False,
            'italic': p.runs[0].italic if p.runs else False,
            'underline': p.runs[0].underline if p.runs else False,
        }

        texto_final = texto_original
        for key in placeholders_no_paragrafo:
            texto_final = texto_final.replace(key, str(contexto[key]))
        
        p.clear()

        texto_restante = texto_final
        for key in placeholders_no_paragrafo:
            valor_placeholder = str(contexto[key])
            partes = texto_restante.split(valor_placeholder, 1)
            
            if partes[0]:
                run_rotulo = aplicar_formatacao_padrao(p.add_run(partes[0]))
                run_rotulo.font.bold = estilo_rotulo['bold']
                run_rotulo.font.italic = estilo_rotulo['italic']
                run_rotulo.underline = estilo_rotulo['underline']

            run_valor = aplicar_formatacao_padrao(p.add_run(valor_placeholder))
            run_valor.font.bold = False
            run_valor.font.italic = False
            run_valor.font.underline = False
            
            texto_restante = partes[1] if len(partes) > 1 else ""

        if texto_restante:
            run_final = aplicar_formatacao_padrao(p.add_run(texto_restante))
            run_final.font.bold = estilo_rotulo['bold']
            run_final.font.italic = estilo_rotulo['italic']
            run_final.underline = estilo_rotulo['underline']

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    processar_paragrafo(p)
    for p in doc.paragraphs:
        processar_paragrafo(p)

def gerar_os(funcionario, riscos_pgr_por_categoria, epis_manuais, medicoes_manuais, riscos_manuais, modelo_doc):
    """
    CORREÇÃO PRINCIPAL: Agora recebe riscos_pgr_por_categoria (dict) ao invés de lista simples
    """
    doc = Document(modelo_doc)
    
    riscos_por_categoria = {cat: [] for cat in CATEGORIAS_RISCO.keys()}
    danos_por_categoria = {cat: [] for cat in CATEGORIAS_RISCO.keys()}

    # CORREÇÃO: Processar riscos PGR por categoria
    for categoria, riscos_lista in riscos_pgr_por_categoria.items():
        if riscos_lista and categoria in RISCOS_PGR_DADOS:
            riscos_por_categoria[categoria].extend(riscos_lista)
            danos = get_danos_por_riscos_pgr(categoria, riscos_lista)
            if danos:
                danos_por_categoria[categoria].append(danos)

    # Processar riscos manuais
    if riscos_manuais:
        map_categorias_rev = {v: k for k, v in CATEGORIAS_RISCO.items()}
        for risco_manual in riscos_manuais:
            categoria_display = risco_manual.get('category')
            categoria_alvo = map_categorias_rev.get(categoria_display)
            if categoria_alvo:
                riscos_por_categoria[categoria_alvo].append(risco_manual.get('risk_name', ''))
                if risco_manual.get('possible_damages'):
                    danos_por_categoria[categoria_alvo].append(risco_manual.get('possible_damages'))

    # Limpar duplicatas
    for cat in danos_por_categoria:
        danos_por_categoria[cat] = sorted(list(set(danos_por_categoria[cat])))

    # Formatar medições
    medicoes_formatadas = []
    for med in medicoes_manuais:
        agente = str(med.get('agent', '')).strip()
        valor = str(med.get('value', '')).strip()
        unidade = str(med.get('unit', '')).strip()
        epi_associado = str(med.get('epi_name', med.get('epi', ''))).strip()
       
        if agente and agente not in ['', 'N/A', 'nan', 'None'] and valor and valor not in ['', 'N/A', 'nan', 'None']:
            linha = f"{agente}: {valor}"
            if unidade and unidade not in ['', 'N/A', 'nan', 'None']:
                linha += f" {unidade}"
            if epi_associado and epi_associado not in ['', 'N/A', 'nan', 'None']:
                linha += f" | EPI: {epi_associado}"
            medicoes_formatadas.append(linha)
    medicoes_texto = "\n".join(medicoes_formatadas) if medicoes_formatadas else "Não aplicável"

    # Processar data de admissão
    data_admissao = "Não informado"
    if 'data_de_admissao' in funcionario and pd.notna(funcionario['data_de_admissao']):
        try: 
            data_admissao = pd.to_datetime(funcionario['data_de_admissao']).strftime('%d/%m/%Y')
        except: 
            data_admissao = str(funcionario['data_de_admissao'])

    # Processar descrição de atividades
    descricao_atividades = "Não informado"
    if 'descricao_de_atividades' in funcionario and pd.notna(funcionario['descricao_de_atividades']):
        descricao_atividades = str(funcionario['descricao_de_atividades']).strip()

    if descricao_atividades in ["Não informado", "", "nan"]:
        funcao = str(funcionario.get('funcao', 'N/A'))
        setor = str(funcionario.get('setor', 'N/A'))
        if funcao != 'N/A' and setor != 'N/A':
            descricao_atividades = f"Atividades relacionadas à função de {funcao} no setor {setor}."
        else:
            descricao_atividades = "Atividades conforme definido pela chefia imediata."

    def tratar_lista_vazia(lista, separador=", "):
        if not lista or all(not item.strip() for item in lista): 
            return "Não identificado"
        return separador.join(sorted(list(set(item for item in lista if item and item.strip()))))

    # Contexto
    contexto = {
        "[NOME EMPRESA]": str(funcionario.get("empresa", "N/A")), 
        "[UNIDADE]": str(funcionario.get("unidade", "N/A")),
        "[NOME FUNCIONÁRIO]": str(funcionario.get("nome_do_funcionario", "N/A")), 
        "[DATA DE ADMISSÃO]": data_admissao,
        "[SETOR]": str(funcionario.get("setor", "N/A")), 
        "[FUNÇÃO]": str(funcionario.get("funcao", "N/A")),
        "[DESCRIÇÃO DE ATIVIDADES]": descricao_atividades,
        "[RISCOS FÍSICOS]": tratar_lista_vazia(riscos_por_categoria["fisico"]),
        "[RISCOS DE ACIDENTE]": tratar_lista_vazia(riscos_por_categoria["acidente"]),
        "[RISCOS QUÍMICOS]": tratar_lista_vazia(riscos_por_categoria["quimico"]),
        "[RISCOS BIOLÓGICOS]": tratar_lista_vazia(riscos_por_categoria["biologico"]),
        "[RISCOS ERGONÔMICOS]": tratar_lista_vazia(riscos_por_categoria["ergonomico"]),
        "[POSSÍVEIS DANOS RISCOS FÍSICOS]": tratar_lista_vazia(danos_por_categoria["fisico"], "; "),
        "[POSSÍVEIS DANOS RISCOS ACIDENTE]": tratar_lista_vazia(danos_por_categoria["acidente"], "; "),
        "[POSSÍVEIS DANOS RISCOS QUÍMICOS]": tratar_lista_vazia(danos_por_categoria["quimico"], "; "),
        "[POSSÍVEIS DANOS RISCOS BIOLÓGICOS]": tratar_lista_vazia(danos_por_categoria["biologico"], "; "),
        "[POSSÍVEIS DANOS RISCOS ERGONÔMICOS]": tratar_lista_vazia(danos_por_categoria["ergonomico"], "; "),
        "[EPIS]": tratar_lista_vazia([epi['epi_name'] for epi in epis_manuais]),
        "[MEDIÇÕES]": medicoes_texto,
    }

    substituir_placeholders(doc, contexto)
    return doc

def main():
    check_authentication()
    init_user_session_state()
    
    if not st.session_state.get('authenticated'):
        show_login_page()
        return
    
    user_id = st.session_state.user_data['user_id']
    show_user_info()
    
    st.markdown('<div class="main-header"><h1>📋 Gerador de Ordens de Serviço</h1><p>Sistema Profissional de Gestão de OS</p></div>', unsafe_allow_html=True)

    # Upload de documentos
    with st.container(border=True):
        st.markdown("#### 📂 1. Documentos Base")
        col1, col2 = st.columns(2)
        with col1:
            arquivo_funcionarios = st.file_uploader("📊 Planilha de Funcionários (.xlsx)", type="xlsx")
        with col2:
            arquivo_modelo_os = st.file_uploader("📄 Modelo de OS (.docx)", type="docx")

    if not
