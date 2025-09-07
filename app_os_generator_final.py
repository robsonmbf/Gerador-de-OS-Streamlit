import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
import zipfile
from io import BytesIO
import time
import re

# --- Configura√ß√£o da P√°gina ---
st.set_page_config(
    page_title="Gerador de Ordens de Servi√ßo (OS)",
    page_icon="üìÑ",
    layout="wide",
)

# --- CSS PERSONALIZADO ---
st.markdown("""
<style>
    [data-testid="stSidebar"] {
        display: none;
    }
    .main-header {
        text-align: center;
        padding-bottom: 20px;
    }
</style>
""", unsafe_allow_html=True)

# --- INICIALIZA√á√ÉO DO SESSION STATE ---
# Usado para manter o estado da aplica√ß√£o entre as intera√ß√µes do usu√°rio.
if 'medicoes_adicionadas' not in st.session_state:
    st.session_state.medicoes_adicionadas = []
if 'riscos_manuais_adicionados' not in st.session_state:
    st.session_state.riscos_manuais_adicionados = []
if 'epis_adicionados' not in st.session_state:
    st.session_state.epis_adicionados = []
if 'setores_concluidos' not in st.session_state:
    st.session_state.setores_concluidos = set()
if 'cargos_concluidos' not in st.session_state:
    st.session_state.cargos_concluidos = set()

# --- LISTAS DE DADOS CONSTANTES ---
UNIDADES_DE_MEDIDA = ["dB(A)", "m/s¬≤", "ppm", "mg/m¬≥", "%", "¬∞C", "lx", "cal/cm¬≤", "¬µT", "kV/m", "W/m¬≤", "f/cm¬≥", "N√£o aplic√°vel"]
AGENTES_DE_RISCO = sorted([
    "Ru√≠do (Cont√≠nuo ou Intermitente)", "Ru√≠do (Impacto)", "Vibra√ß√£o de Corpo Inteiro", "Vibra√ß√£o de M√£os e Bra√ßos",
    "Radia√ß√µes Ionizantes", "Radia√ß√µes N√£o-Ionizantes", "Frio", "Calor", "Press√µes Anormais", "Umidade", "Poeiras", 
    "Fumos", "N√©voas", "Neblinas", "Gases", "Vapores", "Produtos Qu√≠micos em Geral", "V√≠rus", "Bact√©rias", 
    "Protozo√°rios", "Fungos", "Parasitas", "Bacilos"
])
CATEGORIAS_RISCO = {'fisico': 'üî• F√≠sicos', 'quimico': '‚öóÔ∏è Qu√≠micos', 'biologico': 'ü¶† Biol√≥gicos', 'ergonomico': 'üèÉ Ergon√¥micos', 'acidente': '‚ö†Ô∏è Acidentes'}

# --- Fun√ß√µes de L√≥gica de Neg√≥cio ---

def normalizar_texto(texto):
    """Remove acentos, espa√ßos e caracteres especiais para compara√ß√£o de strings."""
    if not isinstance(texto, str): return ""
    texto = texto.lower().strip()
    texto = re.sub(r'[\s\W_]+', '', texto) 
    return texto

def mapear_e_renomear_colunas_funcionarios(df):
    """Renomeia as colunas da planilha de funcion√°rios para um padr√£o conhecido."""
    df_copia = df.copy()
    mapeamento = {
        'nome_do_funcionario': ['nomedofuncionario', 'nome', 'funcionario', 'funcion√°rio', 'colaborador', 'nomecompleto'],
        'funcao': ['funcao', 'fun√ß√£o', 'cargo'],
        'data_de_admissao': ['datadeadmissao', 'dataadmissao', 'admissao', 'admiss√£o'],
        'setor': ['setordetrabalho', 'setor', 'area', '√°rea', 'departamento'],
        'descricao_de_atividades': ['descricaodeatividades', 'atividades', 'descricaoatividades', 'descri√ß√£odeatividades', 'tarefas', 'descricaodastarefas'],
        'empresa': ['empresa'],
        'unidade': ['unidade']
    }
    
    colunas_renomeadas = {}
    colunas_df_normalizadas = {normalizar_texto(col): col for col in df_copia.columns}
    for nome_padrao, nomes_possiveis in mapeamento.items():
        for nome_possivel in nomes_possiveis:
            if nome_possivel in colunas_df_normalizadas:
                coluna_original = colunas_df_normalizadas[nome_possivel]
                colunas_renomeadas[coluna_original] = nome_padrao
                break
    df_copia.rename(columns=colunas_renomeadas, inplace=True)
    return df_copia

@st.cache_data
def carregar_planilha(arquivo):
    """Carrega e armazena em cache a planilha para evitar recarregamentos."""
    if arquivo is None: return None
    try:
        return pd.read_excel(arquivo)
    except Exception as e:
        st.error(f"Erro ao ler o ficheiro Excel: {e}")
        return None

def obter_dados_pgr():
    """Simula a obten√ß√£o de dados de um PGR. Em um caso real, isso viria de um banco de dados ou outra planilha."""
    data = [
        # Riscos F√≠sicos
        {'categoria': 'fisico', 'risco': 'Ru√≠do (Cont√≠nuo ou Intermitente)', 'possiveis_danos': 'Perda auditiva, zumbido, estresse, irritabilidade.'},
        {'categoria': 'fisico', 'risco': 'Ru√≠do (Impacto)', 'possiveis_danos': 'Perda auditiva, trauma ac√∫stico.'},
        {'categoria': 'fisico', 'risco': 'Vibra√ß√£o de Corpo Inteiro', 'possiveis_danos': 'Problemas na coluna, dores lombares.'},
        {'categoria': 'fisico', 'risco': 'Vibra√ß√£o de M√£os e Bra√ßos', 'possiveis_danos': 'Doen√ßas osteomusculares, problemas circulat√≥rios (s√≠ndrome de Raynaud).'},
        {'categoria': 'fisico', 'risco': 'Calor', 'possiveis_danos': 'Desidrata√ß√£o, insola√ß√£o, c√£ibras, exaust√£o, interma√ß√£o.'},
        {'categoria': 'fisico', 'risco': 'Frio', 'possiveis_danos': 'Hipotermia, congelamento, doen√ßas respirat√≥rias.'},
        {'categoria': 'fisico', 'risco': 'Radia√ß√µes Ionizantes', 'possiveis_danos': 'C√¢ncer, muta√ß√µes gen√©ticas, queimaduras.'},
        {'categoria': 'fisico', 'risco': 'Radia√ß√µes N√£o-Ionizantes', 'possiveis_danos': 'Queimaduras, les√µes oculares, c√¢ncer de pele.'},
        {'categoria': 'fisico', 'risco': 'Press√µes Anormais', 'possiveis_danos': 'Doen√ßa descompressiva, barotrauma.'},
        {'categoria': 'fisico', 'risco': 'Umidade', 'possiveis_danos': 'Doen√ßas respirat√≥rias, dermatites, micoses.'},

        # Riscos Qu√≠micos
        {'categoria': 'quimico', 'risco': 'Poeiras', 'possiveis_danos': 'Pneumoconioses (silicose, asbestose), irrita√ß√£o respirat√≥ria, alergias.'},
        {'categoria': 'quimico', 'risco': 'Fumos', 'possiveis_danos': 'Doen√ßas respirat√≥rias (febre dos fumos met√°licos), intoxica√ß√µes.'},
        {'categoria': 'quimico', 'risco': 'N√©voas', 'possiveis_danos': 'Irrita√ß√£o respirat√≥ria, dermatites.'},
        {'categoria': 'quimico', 'risco': 'Neblinas', 'possiveis_danos': 'Irrita√ß√£o do trato respirat√≥rio.'},
        {'categoria': 'quimico', 'risco': 'Gases', 'possiveis_danos': 'Asfixia, intoxica√ß√µes, irrita√ß√£o respirat√≥ria.'},
        {'categoria': 'quimico', 'risco': 'Vapores', 'possiveis_danos': 'Irrita√ß√£o respirat√≥ria, intoxica√ß√µes, dermatites.'},
        {'categoria': 'quimico', 'risco': 'Produtos Qu√≠micos em Geral', 'possiveis_danos': 'Queimaduras, irrita√ß√µes, intoxica√ß√µes, dermatites, c√¢ncer.'},

        # Riscos Biol√≥gicos
        {'categoria': 'biologico', 'risco': 'Bact√©rias', 'possiveis_danos': 'Infec√ß√µes, doen√ßas infecciosas (t√©tano, tuberculose).'},
        {'categoria': 'biologico', 'risco': 'Fungos', 'possiveis_danos': 'Micoses, alergias, infec√ß√µes respirat√≥rias.'},
        {'categoria': 'biologico', 'risco': 'V√≠rus', 'possiveis_danos': 'Doen√ßas virais (hepatite, HIV), infec√ß√µes.'},
        {'categoria': 'biologico', 'risco': 'Parasitas', 'possiveis_danos': 'Doen√ßas parasit√°rias, infec√ß√µes.'},
        {'categoria': 'biologico', 'risco': 'Protozo√°rios', 'possiveis_danos': 'Doen√ßas parasit√°rias (leishmaniose, mal√°ria).'},
        {'categoria': 'biologico', 'risco': 'Bacilos', 'possiveis_danos': 'Infec√ß√µes diversas, como tuberculose.'},
        
        # Riscos Ergon√¥micos
        {'categoria': 'ergonomico', 'risco': 'Levantamento e Transporte Manual de Peso', 'possiveis_danos': 'Les√µes musculoesquel√©ticas, dores na coluna.'},
        {'categoria': 'ergonomico', 'risco': 'Posturas Inadequadas', 'possiveis_danos': 'Dores musculares, les√µes na coluna, LER/DORT.'},
        {'categoria': 'ergonomico', 'risco': 'Repetitividade', 'possiveis_danos': 'LER/DORT, tendinites, s√≠ndrome do t√∫nel do carpo.'},
        {'categoria': 'ergonomico', 'risco': 'Jornada de Trabalho Prolongada', 'possiveis_danos': 'Fadiga, estresse, acidentes de trabalho.'},
        {'categoria': 'ergonomico', 'risco': 'Ilumina√ß√£o Inadequada', 'possiveis_danos': 'Fadiga visual, dores de cabe√ßa, acidentes.'},

        # Riscos de Acidentes
        {'categoria': 'acidente', 'risco': 'Arranjo F√≠sico Inadequado', 'possiveis_danos': 'Quedas, colis√µes, esmagamentos.'},
        {'categoria': 'acidente', 'risco': 'M√°quinas e Equipamentos sem Prote√ß√£o', 'possiveis_danos': 'Amputa√ß√µes, cortes, esmagamentos, prensamentos.'},
        {'categoria': 'acidente', 'risco': 'Ferramentas Inadequadas ou Defeituosas', 'possiveis_danos': 'Cortes, perfura√ß√µes, fraturas.'},
        {'categoria': 'acidente', 'risco': 'Eletricidade', 'possiveis_danos': 'Choque el√©trico, queimaduras, fibrila√ß√£o ventricular.'},
        {'categoria': 'acidente', 'risco': 'Inc√™ndio e Explos√£o', 'possiveis_danos': 'Queimaduras, asfixia, les√µes por impacto.'},
        {'categoria': 'acidente', 'risco': 'Animais Pe√ßonhentos', 'possiveis_danos': 'Picadas, mordidas, rea√ß√µes al√©rgicas, envenenamento.'},
        {'categoria': 'acidente', 'risco': 'Armazenamento Inadequado', 'possiveis_danos': 'Quedas de materiais, esmagamentos.'},
        {'categoria': 'acidente', 'risco': 'Trabalho em Altura', 'possiveis_danos': 'Quedas, fraturas, morte.'},
        {'categoria': 'acidente', 'risco': 'Espa√ßos Confinados', 'possiveis_danos': 'Asfixia, intoxica√ß√µes, explos√µes.'},
        {'categoria': 'acidente', 'risco': 'Proje√ß√£o de Part√≠culas', 'possiveis_danos': 'Les√µes oculares, cortes na pele.'}
    ]
    return pd.DataFrame(data)


def substituir_placeholders(doc, contexto):
    """Substitui chaves de texto (ex: [NOME]) no documento Word pelos valores do contexto."""
    # Itera por par√°grafos, tabelas e cabe√ßalhos para uma substitui√ß√£o completa.
    all_elements = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_elements.extend(cell.paragraphs)
    for section in doc.sections:
        header = section.header
        all_elements.extend(header.paragraphs)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_elements.extend(cell.paragraphs)

    for p in all_elements:
        full_text = "".join(run.text for run in p.runs)
        if not full_text.strip(): continue

        # Otimiza a substitui√ß√£o para evitar corromper a formata√ß√£o
        for key, value in contexto.items():
            if key in full_text:
                full_text = full_text.replace(str(key), str(value))
        
        # Limpa o par√°grafo e adiciona o novo texto com formata√ß√£o padr√£o
        for i in range(len(p.runs)):
            p.runs[i].text = ''
        if p.runs:
            run = p.runs[0]
            run.text = full_text
            font = run.font
            font.name = 'Segoe UI'
            font.size = Pt(9)
            
def gerar_os(funcionario, df_pgr, riscos_selecionados, epis_manuais, medicoes_manuais, riscos_manuais, modelo_doc_carregado):
    """Fun√ß√£o principal que gera um documento de Ordem de Servi√ßo para um funcion√°rio."""
    doc = Document(modelo_doc_carregado)
    
    # Processa riscos do PGR
    riscos_info = df_pgr[df_pgr['risco'].isin(riscos_selecionados)]
    riscos_por_categoria = {cat: [] for cat in CATEGORIAS_RISCO.keys()}
    danos_por_categoria = {cat: [] for cat in CATEGORIAS_RISCO.keys()}
    for _, risco_row in riscos_info.iterrows():
        categoria = str(risco_row.get("categoria", "")).lower()
        if categoria in riscos_por_categoria:
            riscos_por_categoria[categoria].append(str(risco_row.get("risco", "")))
            danos = risco_row.get("possiveis_danos")
            if pd.notna(danos): danos_por_categoria[categoria].append(str(danos))

    # Adiciona riscos inseridos manualmente
    if riscos_manuais:
        map_categorias_rev = {v: k for k, v in CATEGORIAS_RISCO.items()}
        for risco_manual in riscos_manuais:
            categoria_display = risco_manual.get('categoria')
            categoria_alvo = map_categorias_rev.get(categoria_display)
            if categoria_alvo:
                riscos_por_categoria[categoria_alvo].append(risco_manual.get('risco', ''))
                if risco_manual.get('danos'):
                    danos_por_categoria[categoria_alvo].append(risco_manual.get('danos'))

    # Formata a lista de medi√ß√µes, incluindo o EPI associado
    medicoes_formatadas = []
    for med in medicoes_manuais:
        epi_info = f" | EPI: {med['epi']}" if med.get("epi") else ""
        medicoes_formatadas.append(f"{med['agente']}: {med['valor']} {med['unidade']}{epi_info}")
    medicoes_texto = "\n".join(medicoes_formatadas) if medicoes_formatadas else "N√£o aplic√°vel"

    # Trata e formata campos do funcion√°rio
    data_admissao = "N√£o informado"
    if 'data_de_admissao' in funcionario and pd.notna(funcionario['data_de_admissao']):
        try: data_admissao = pd.to_datetime(funcionario['data_de_admissao']).strftime('%d/%m/%Y')
        except Exception: data_admissao = str(funcionario['data_de_admissao'])

    descricao_atividades = "N√£o informado"
    if 'descricao_de_atividades' in funcionario and pd.notna(funcionario['descricao_de_atividades']):
        descricao_atividades = str(funcionario['descricao_de_atividades'])
    
    def tratar_lista_vazia(lista, separador=", "):
        if not lista or all(not item.strip() for item in lista): return "N√£o identificado"
        return separador.join(sorted(list(set(item for item in lista if item and item.strip()))))

    # Cria o dicion√°rio de contexto para substitui√ß√£o no Word
    contexto = {
        "[NOME EMPRESA]": str(funcionario.get("empresa", "N/A")), 
        "[UNIDADE]": str(funcionario.get("unidade", "N/A")),
        "[NOME FUNCION√ÅRIO]": str(funcionario.get("nome_do_funcionario", "N/A")), 
        "[DATA DE ADMISS√ÉO]": data_admissao,
        "[SETOR]": str(funcionario.get("setor", "N/A")), 
        "[FUN√á√ÉO]": str(funcionario.get("funcao", "N/A")),
        "[DESCRI√á√ÉO DE ATIVIDADES]": descricao_atividades,
        "[RISCOS F√çSICOS]": tratar_lista_vazia(riscos_por_categoria["fisico"]),
        "[RISCOS DE ACIDENTE]": tratar_lista_vazia(riscos_por_categoria["acidente"]),
        "[RISCOS QU√çMICOS]": tratar_lista_vazia(riscos_por_categoria["quimico"]),
        "[RISCOS BIOL√ìGICOS]": tratar_lista_vazia(riscos_por_categoria["biologico"]),
        "[RISCOS ERGON√îMICOS]": tratar_lista_vazia(riscos_por_categoria["ergonomico"]),
        "[POSS√çVEIS DANOS RISCOS F√çSICOS]": tratar_lista_vazia(danos_por_categoria["fisico"], "; "),
        "[POSS√çVEIS DANOS RISCOS ACIDENTE]": tratar_lista_vazia(danos_por_categoria["acidente"], "; "),
        "[POSS√çVEIS DANOS RISCOS QU√çMICOS]": tratar_lista_vazia(danos_por_categoria["quimico"], "; "),
        "[POSS√çVEIS DANOS RISCOS BIOL√ìGICOS]": tratar_lista_vazia(danos_por_categoria["biologico"], "; "),
        "[POSS√çVEIS DANOS RISCOS ERGON√îMICOS]": tratar_lista_vazia(danos_por_categoria["ergonomico"], "; "),
        "[EPIS]": tratar_lista_vazia(epis_manuais.split(',')) or "N√£o aplic√°vel",
        "[MEDI√á√ïES]": medicoes_texto,
    }
    
    substituir_placeholders(doc, contexto)
    return doc

# --- Interface Gr√°fica do Streamlit ---
st.markdown("""<div class="main-header"><h1>üìÑ Gerador de Ordens de Servi√ßo (OS)</h1><p>Gere OS em lote a partir de um modelo Word (.docx) e uma planilha de funcion√°rios.</p></div>""", unsafe_allow_html=True)

with st.container(border=True):
    st.markdown("##### üìÇ 1. Carregue os Documentos")
    col1, col2 = st.columns(2)
    with col1:
        arquivo_funcionarios = st.file_uploader("üìÑ **Planilha de Funcion√°rios (.xlsx)**", type="xlsx", help="Planilha com colunas como: Nome, Fun√ß√£o, Setor, Empresa, etc.")
    with col2:
        arquivo_modelo_os = st.file_uploader("üìù **Modelo de OS (.docx)**", type="docx", help="Documento Word com placeholders como [NOME FUNCION√ÅRIO], [SETOR], etc.")

if not arquivo_funcionarios or not arquivo_modelo_os:
    st.info("üìã Por favor, carregue a Planilha de Funcion√°rios e o Modelo de OS para continuar.")
    st.stop()

# Carrega os dados ap√≥s os uploads
df_funcionarios_raw = carregar_planilha(arquivo_funcionarios)
df_funcionarios = mapear_e_renomear_colunas_funcionarios(df_funcionarios_raw)
df_pgr = obter_dados_pgr()

# --- Se√ß√£o de Filtros e Sele√ß√£o ---
with st.container(border=True):
    st.markdown('##### üë• 2. Selecione os Funcion√°rios')
    
    col_filtro1, col_filtro2 = st.columns([3,1])
    with col_filtro1:
        setores = sorted(df_funcionarios['setor'].dropna().unique().tolist()) if 'setor' in df_funcionarios.columns else []
        setor_sel = st.multiselect("Filtrar por Setor(es)", setores, placeholder="Selecione um ou mais setores")
    
    df_filtrado_setor = df_funcionarios[df_funcionarios['setor'].isin(setor_sel)] if setor_sel else df_funcionarios
    
    with col_filtro1:
        funcoes = sorted(df_filtrado_setor['funcao'].dropna().unique().tolist()) if 'funcao' in df_filtrado_setor.columns else []
        funcao_sel = st.multiselect("Filtrar por Fun√ß√£o/Cargo(s)", funcoes, placeholder="Selecione uma ou mais fun√ß√µes")
    
    with col_filtro2:
        st.write("") # Espa√ßamento
        st.write("") # Espa√ßamento
        if st.button("Limpar Filtros"):
            st.rerun()

    df_final_filtrado = df_filtrado_setor[df_filtrado_setor['funcao'].isin(funcao_sel)] if funcao_sel else df_filtrado_setor
    
    st.success(f"**{len(df_final_filtrado)} funcion√°rios selecionados.** Revise a lista abaixo.")
    
    colunas_desejadas = ['nome_do_funcionario', 'setor', 'funcao']
    colunas_existentes = [col for col in colunas_desejadas if col in df_final_filtrado.columns]
    if not colunas_existentes:
        st.error("‚ùå Nenhuma das colunas essenciais (nome, setor, fun√ß√£o) foi encontrada na planilha.")
    else:
        st.dataframe(df_final_filtrado[colunas_existentes], use_container_width=True, height=250)

# --- Se√ß√£o de Configura√ß√£o de Riscos e Medidas ---
with st.container(border=True):
    st.markdown('##### ‚ö†Ô∏è 3. Configure os Riscos e Medidas de Controle')
    
    st.info("Os riscos, medi√ß√µes e EPIs configurados aqui ser√£o aplicados a **TODOS** os funcion√°rios selecionados acima.")

    # --- Sele√ß√£o de Riscos do PGR ---
    st.markdown("**Riscos Identificados (PGR)**")
    riscos_selecionados = []
    tabs = st.tabs(list(CATEGORIAS_RISCO.values()))
    for i, (key, nome) in enumerate(CATEGORIAS_RISCO.items()):
        with tabs[i]:
            riscos_categoria = df_pgr[df_pgr['categoria'] == key]['risco'].tolist()
            selecionados = st.multiselect(f"Selecione os riscos:", options=riscos_categoria, key=f"riscos_{key}", default=[])
            riscos_selecionados.extend(selecionados)

    # --- Resumo dos Riscos Selecionados ---
    if riscos_selecionados:
        with st.expander(f"**Resumo de Riscos Selecionados ({len(riscos_selecionados)} no total)**", expanded=True):
            riscos_categorizados_para_display = {}
            for risco_nome in sorted(riscos_selecionados):
                categoria_key_series = df_pgr[df_pgr['risco'] == risco_nome]['categoria']
                if not categoria_key_series.empty:
                    categoria_key = categoria_key_series.iloc[0]
                    categoria_display = CATEGORIAS_RISCO.get(categoria_key, "Outros")
                    if categoria_display not in riscos_categorizados_para_display:
                        riscos_categorizados_para_display[categoria_display] = []
                    riscos_categorizados_para_display[categoria_display].append(risco_nome)
            
            for categoria, lista_riscos in riscos_categorizados_para_display.items():
                st.markdown(f"**{categoria}**")
                for risk in lista_riscos:
                    st.markdown(f"&nbsp;&nbsp;&nbsp; - {risk}")

    # --- Adi√ß√£o de Medi√ß√µes, Riscos Manuais e EPIs ---
    col_exp1, col_exp2, col_exp3 = st.columns(3)
    with col_exp1:
        with st.expander("üìä **Adicionar Medi√ß√µes**"):
            def adicionar_medicao():
                agente_sel = st.session_state.agente_input
                agente_man = st.session_state.get("agente_manual_input")
                agente_final = agente_man if agente_sel == "Outro (digitar manualmente)" else agente_sel
                valor = st.session_state.valor_input
                unidade = st.session_state.unidade_input
                epi = st.session_state.epi_medicao_input

                if agente_final and valor:
                    medicao = {"agente": agente_final, "valor": valor, "unidade": unidade, "epi": epi}
                    st.session_state.medicoes_adicionadas.append(medicao)
                    # Limpa os campos para a pr√≥xima entrada
                    st.session_state.agente_input = AGENTES_DE_RISCO[0]
                    st.session_state.valor_input = ""
                    st.session_state.epi_medicao_input = ""
                    if "agente_manual_input" in st.session_state: st.session_state.agente_manual_input = ""
                else: st.warning("Preencha o Agente e o Valor.")
            
            agente_opts = [""] + AGENTES_DE_RISCO + ["Outro (digitar manualmente)"]
            agente_selecionado = st.selectbox("Agente/Fonte", options=agente_opts, key="agente_input")
            if agente_selecionado == "Outro (digitar manualmente)":
                st.text_input("Digite o Agente", key="agente_manual_input")
            st.text_input("Valor Medido", key="valor_input")
            st.selectbox("Unidade", UNIDADES_DE_MEDIDA, key="unidade_input")
            st.text_input("EPI Associado", key="epi_medicao_input", help="EPI recomendado para esta medi√ß√£o espec√≠fica.")
            st.button("Adicionar Medi√ß√£o", on_click=adicionar_medicao, use_container_width=True)

            if st.session_state.medicoes_adicionadas:
                st.write("**Adicionadas:**")
                for i, med in enumerate(st.session_state.medicoes_adicionadas):
                    epi_info = f" | EPI: {med['epi']}" if med.get("epi") else ""
                    st.markdown(f"- {med['agente']}: {med['valor']} {med['unidade']}{epi_info}")
                if st.button("Limpar Medi√ß√µes", key="limpar_med", use_container_width=True):
                    st.session_state.medicoes_adicionadas = []
                    st.rerun()

    with col_exp2:
        with st.expander("‚ûï **Adicionar Risco Manual**"):
            def adicionar_risco_manual():
                if st.session_state.risco_input and st.session_state.categoria_input:
                    risco = {"risco": st.session_state.risco_input, "categoria": st.session_state.categoria_input, "danos": st.session_state.danos_input}
                    st.session_state.riscos_manuais_adicionados.append(risco)
                    st.session_state.risco_input = ""; st.session_state.danos_input = ""
                else: st.warning("Preencha a Descri√ß√£o e Categoria.")
            
            st.text_input("Descri√ß√£o do Risco", key="risco_input")
            st.selectbox("Categoria", [""] + list(CATEGORIAS_RISCO.values()), key="categoria_input")
            st.text_area("Poss√≠veis Danos", key="danos_input", height=50)
            st.button("Adicionar Risco", on_click=adicionar_risco_manual, use_container_width=True)

            if st.session_state.riscos_manuais_adicionados:
                st.write("**Adicionados:**")
                for r in st.session_state.riscos_manuais_adicionados:
                    st.markdown(f"- **{r['risco']}** ({r['categoria']})")
                if st.button("Limpar Riscos", key="limpar_ris", use_container_width=True):
                    st.session_state.riscos_manuais_adicionados = []
                    st.rerun()

    with col_exp3:
        with st.expander("ü¶∫ **Adicionar EPIs Gerais**"):
            def adicionar_epi():
                if st.session_state.epi_input.strip():
                    st.session_state.epis_adicionados.append(st.session_state.epi_input.strip())
                    st.session_state.epi_input = ""

            st.text_input("Nome do EPI", key="epi_input", help="EPIs de uso geral, n√£o ligados a uma medi√ß√£o espec√≠fica.")
            st.button("Adicionar EPI", on_click=adicionar_epi, use_container_width=True)
            
            if st.session_state.epis_adicionados:
                st.write("**Adicionados:**")
                for epi_item in st.session_state.epis_adicionados:
                    st.markdown(f"- {epi_item}")
                if st.button("Limpar EPIs", key="limpar_epi", use_container_width=True):
                    st.session_state.epis_adicionados = []
                    st.rerun()

# --- Se√ß√£o de Gera√ß√£o dos Documentos ---
st.divider()
if st.button("üöÄ Gerar OS para Funcion√°rios Selecionados", type="primary", use_container_width=True, disabled=df_final_filtrado.empty):
    epis_finais = ", ".join(st.session_state.epis_adicionados)
    
    with st.spinner(f"Gerando {len(df_final_filtrado)} documentos... Por favor, aguarde."):
        documentos_gerados = []
        os_geradas_info_batch = [] 
        
        progresso = st.progress(0)
        total_funcs = len(df_final_filtrado)

        for i, (_, func) in enumerate(df_final_filtrado.iterrows()):
            doc = gerar_os(
                func, 
                df_pgr, 
                riscos_selecionados, 
                epis_finais, 
                st.session_state.medicoes_adicionadas, 
                st.session_state.riscos_manuais_adicionados, 
                arquivo_modelo_os
            )
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            nome_limpo = re.sub(r'[^\w\s-]', '', func.get("nome_do_funcionario", "Func_Sem_Nome")).strip().replace(" ", "_")
            documentos_gerados.append((f"OS_{nome_limpo}.docx", doc_io.getvalue()))
            os_geradas_info_batch.append({
                'Funcion√°rio': func.get("nome_do_funcionario", "N/A"),
                'Setor': func.get("setor", "N/A"),
                'Cargo/Fun√ß√£o': func.get("funcao", "N/A")
            })
            progresso.progress((i + 1) / total_funcs)

    if documentos_gerados:
        df_resumo_batch = pd.DataFrame(os_geradas_info_batch)
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for nome_arquivo, conteudo_doc in documentos_gerados:
                zip_file.writestr(nome_arquivo, conteudo_doc)
        
        setores_nome = "_".join(set(s.replace(" ", "") for s in setor_sel)) if setor_sel else "TODOS"
        funcoes_nome = "_".join(set(f.replace(" ", "") for f in funcao_sel)) if funcao_sel else "TODAS"
        nome_arquivo_zip = f"OS_{setores_nome}_{funcoes_nome}_{time.strftime('%Y%m%d')}.zip"
        
        st.success(f"üéâ **{len(documentos_gerados)} Ordens de Servi√ßo geradas com sucesso!**")
        st.download_button(
            label="üì• Baixar Todas as OS (.zip)", 
            data=zip_buffer.getvalue(), 
            file_name=nome_arquivo_zip, 
            mime="application/zip",
            use_container_width=True
        )
        with st.expander("üìÑ Ver resumo do lote gerado", expanded=True):
            st.dataframe(df_resumo_batch, use_container_width=True)

